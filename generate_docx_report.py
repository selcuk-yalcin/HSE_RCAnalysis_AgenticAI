"""
DOCX RAPOR OLUŞTURUCU - MPT RAMAK KALA ANALIZI
==============================================

JSON sonuçlarından profesyonel DOCX rapor oluştur
"""

import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_full_report_docx(json_file_path, output_path):
    """JSON analiz dosyasından DOCX rapor oluştur."""
    
    # JSON'u oku
    with open(json_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Yeni DOCX document oluştur
    doc = Document()
    
    # ========================================================================
    # BAŞLIK VE KAPAK
    # ========================================================================
    
    title = doc.add_heading('HSE ROOT CAUSE ANALYSIS RAPORU', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('MPT Test Sahası - Sarmal Kapı Düşen Parça Ramak Kala Olayi', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Rapor bilgileri
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_data = [
        ("Referans No:", data.get('incident_ref', 'N/A')),
        ("Rapor Tarihi:", datetime.now().strftime("%d.%m.%Y")),
        ("Analiz Yöntemi:", data.get('analysis_method', 'N/A')),
        ("Analiz Saati:", datetime.now().strftime("%H:%M:%S")),
        ("Olay Tarihi:", "20.01.2026 09:10"),
        ("Olay Yeri:", "MPT Test Sahası")
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)
    
    doc.add_page_break()
    
    # ========================================================================
    # 1. ÖZET
    # ========================================================================
    
    doc.add_heading('1. OLAY ÖZETI', level=1)
    
    overview = data.get('overview', {})
    
    p = doc.add_paragraph()
    p.add_run('Olay Tipi: ').bold = True
    p.add_run(overview.get('incident_type', 'N/A'))
    
    p = doc.add_paragraph()
    p.add_run('Ciddiyet: ').bold = True
    p.add_run(overview.get('severity_level', 'N/A'))
    
    p = doc.add_paragraph()
    p.add_run('Potansiyel Zararlar: ').bold = True
    for harm in overview.get('potential_harm', []):
        doc.add_paragraph(harm, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Özet: ').bold = True
    p.add_run(overview.get('summary', 'N/A'))
    
    doc.add_page_break()
    
    # ========================================================================
    # 2. RİSK DEĞERLENDİRMESİ
    # ========================================================================
    
    doc.add_heading('2. RİSK DEĞERLENDİRMESİ', level=1)
    
    assessment = data.get('assessment', {})
    
    p = doc.add_paragraph()
    p.add_run('Değerlendirme: ').bold = True
    p.add_run(assessment.get('assessment_result', 'N/A'))
    
    p = doc.add_paragraph()
    p.add_run('RIDDOR Rapor Edilebilir: ').bold = True
    p.add_run(assessment.get('riddor_reportable', 'N/A'))
    
    p = doc.add_paragraph()
    p.add_run('Araştırma Seviyesi: ').bold = True
    p.add_run(assessment.get('investigation_level', 'N/A'))
    
    doc.add_paragraph()
    doc.add_heading('Potansiyel Sonuçlar:', level=3)
    for consequence in assessment.get('potential_consequences', []):
        doc.add_paragraph(consequence, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_heading('Ramak Kala Analizi:', level=3)
    p = doc.add_paragraph(assessment.get('near_miss_analysis', 'N/A'))
    
    doc.add_page_break()
    
    # ========================================================================
    # 3. KÖK NEDEN ANALİZİ
    # ========================================================================
    
    doc.add_heading('3. KÖK NEDEN ANALİZİ (5-WHY)', level=1)
    
    rca = data.get('root_cause_analysis', {})
    
    # Meta Root Cause
    meta = rca.get('meta_root_cause', {})
    if meta:
        doc.add_heading('Meta Kök Neden (Ortak Payda):', level=3)
        p = doc.add_paragraph()
        p.add_run(f"[{meta.get('code', '?')}] ").bold = True
        p.add_run(meta.get('name', 'N/A'))
        doc.add_paragraph(meta.get('description', 'N/A'))
    
    # Kök Nedenler
    doc.add_paragraph()
    doc.add_heading('Tespit Edilen Kök Nedenler:', level=3)
    
    for i, rc in enumerate(rca.get('final_root_causes', []), 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. [{rc.get('code', '?')}] ").bold = True
        p.add_run(rc.get('name', 'N/A'))
        doc.add_paragraph(rc.get('description', 'N/A'), style='List Bullet')
    
    # Analiz Dalları
    doc.add_paragraph()
    doc.add_heading('Analiz Dalları (5-WHY Zinciri):', level=3)
    
    for i, branch in enumerate(rca.get('analysis_branches', []), 1):
        doc.add_paragraph()
        p = doc.add_paragraph(f"DAL {i}:")
        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.25)
        
        direct = branch.get('direct_cause', {})
        p = doc.add_paragraph(f"Doğrudan Neden: {direct.get('name', 'N/A')}")
        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.5)
        
        root = branch.get('root_cause', {})
        p = doc.add_paragraph(f"Kök Neden: {root.get('name', 'N/A')}")
        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.5)
        
        doc.add_paragraph()
        p = doc.add_paragraph("5-WHY Zinciri:")
        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.5)
        
        for why in branch.get('five_why_chain', []):
            p = doc.add_paragraph(f"WHY: {why.get('why', 'N/A')}")
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.75)
            
            p = doc.add_paragraph(f"BECAUSE: {why.get('because', 'N/A')}")
            p_format = p.paragraph_format
            p_format.left_indent = Inches(1.0)
    
    doc.add_page_break()
    
    # ========================================================================
    # 4. ACTION PLAN
    # ========================================================================
    
    doc.add_heading('4. DÜZELTICI VE ÖNLEYICI TEDİRLER', level=1)
    
    actions = data.get('action_plan', {})
    
    # Acil Tedbirler
    doc.add_heading('Acil Tedbirler (KRITIK):', level=3)
    
    for action in actions.get('immediate_actions', []):
        p = doc.add_paragraph()
        p.add_run(f"• {action.get('description', 'N/A')} ").bold = True
        p.add_run(f"[{action.get('priority', 'N/A')}, Owner: {action.get('owner', 'N/A')}, Deadline: {action.get('deadline', 'N/A')}]")
    
    # Uzun Vadeli Önleyici Tedbirler
    doc.add_paragraph()
    doc.add_heading('Uzun Vadeli Önleyici Tedbirler:', level=3)
    
    for i, measure in enumerate(actions.get('preventive_measures', []), 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {measure.get('description', 'N/A')} ").bold = True
        p.add_run(f"[Owner: {measure.get('owner', 'N/A')}, Timeline: {measure.get('timeline', 'N/A')}]")
    
    doc.add_page_break()
    
    # ========================================================================
    # 5. BAŞARI KRİTERLERİ
    # ========================================================================
    
    doc.add_heading('5. BAŞARI KRİTERLERİ VE İZLEME', level=1)
    
    doc.add_paragraph("Tedbirlerin başarıyla uygulandığını göstermek için aşağıdaki kriterler karşılanmalıdır:")
    doc.add_paragraph()
    
    success_criteria = [
        "Tüm sarmal kapılar bakım kayıtları güncel (30 gün içinde)",
        "Çalışan eğitimi tamamlandı (%100)",
        "Anormal ses prosedürü yazılı ve afişe edildi",
        "Hiç bir incident/near-miss (6 ay boyunca)",
        "Bakım sistem veri tabanı live ve aktif (tracking)",
        "Güvenlik sensörleri kuruldu ve test edildi",
        "Periyodik denetim sistemi oluşturuldu (3-6 aylık aralıklar)"
    ]
    
    for criteria in success_criteria:
        doc.add_paragraph(criteria, style='List Bullet')
    
    doc.add_page_break()
    
    # ========================================================================
    # SON SAYFA - ONAYLAR
    # ========================================================================
    
    doc.add_heading('6. ONAYLAR', level=1)
    
    doc.add_paragraph()
    doc.add_paragraph("Bu rapor aşağıdaki kişilerce incelenmiş ve onaylanmıştır:")
    doc.add_paragraph()
    
    approval_table = doc.add_table(rows=5, cols=3)
    approval_table.style = 'Light Grid Accent 1'
    
    approval_table.rows[0].cells[0].text = "Rol"
    approval_table.rows[0].cells[1].text = "Ad Soyad"
    approval_table.rows[0].cells[2].text = "Tarih / İmza"
    
    approval_table.rows[1].cells[0].text = "HSE Müdürü"
    approval_table.rows[2].cells[0].text = "Operasyon Müdürü"
    approval_table.rows[3].cells[0].text = "Tesisler Müdürü"
    approval_table.rows[4].cells[0].text = "Bölge Sorumlusu"
    
    doc.add_paragraph()
    doc.add_paragraph(f"Rapor Oluşturulan Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
    
    # DOCX kaydet
    doc.save(output_path)
    
    return output_path


# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    import sys
    
    # JSON dosyasını oku
    json_file = "outputs/mpt_falling_part_near_miss/full_analysis_20260417_185804.json"
    output_file = "outputs/mpt_falling_part_near_miss/MPT_Ramak_Kala_RCA_Raporu.docx"
    
    if not Path(json_file).exists():
        print(f"❌ JSON dosyası bulunamadı: {json_file}")
        sys.exit(1)
    
    print("🤖 DOCX Rapor Oluşturuluyor...")
    print(f"   📄 Giriş: {json_file}")
    
    try:
        result = create_full_report_docx(json_file, output_file)
        
        file_size = Path(result).stat().st_size
        
        print(f"\n✅ DOCX RAPOR BAŞARIYLA OLUŞTURULDU!")
        print(f"   📄 Çıkış: {result}")
        print(f"   📊 Dosya Boyutu: {file_size:,} bytes")
        print(f"\n   Raporun Kapsamı:")
        print(f"      • Kapak ve Özet Bilgiler")
        print(f"      • Olay Özeti ve Ciddiyet")
        print(f"      • Risk Değerlendirmesi")
        print(f"      • Kök Neden Analizi (5-WHY)")
        print(f"      • Düzeltici/Önleyici Tedbirler")
        print(f"      • Başarı Kriterleri")
        print(f"      • Onay Sayfası")
        
        sys.exit(0)
        
    except Exception as e:
        print(f"\n❌ HATA: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
