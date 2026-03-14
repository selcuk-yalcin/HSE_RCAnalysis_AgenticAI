"""
Multilingual Taxonomy Parser - DOCX'ten Çok Dilli JSON'a Dönüştürücü
====================================================================

Bu script, knowlodge_base/ klasöründeki tüm DOCX dosyalarını okur ve
tek bir çok dilli JSON dosyası oluşturur.

Kullanım:
    python rag_pipeline/parsing/parse_multilingual_taxonomy.py
"""

import sys
import json
import re
from pathlib import Path
from typing import List, Dict, Optional
from docx import Document

# Proje kök dizinini Python path'e ekle
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

from rag_pipeline.schemas.cause_models import Cause, LocalizedContent, ExclusionCondition, Taxonomy


class MultilingualTaxonomyParser:
    """Çok dilli DOCX dosyalarından taksonomi verilerini ayrıştırır."""
    
    # Dil tespiti için anahtar kelimeler
    LANGUAGE_MARKERS = {
        'tr': ['Seç eğer:', 'Bu değil eğer:', 'Tipik:', 'İLK GÖRÜNÜR NEDENLER'],
        'en': ['Choose if:', 'Not this if:', 'Typical:', 'IMMEDIATE CAUSES']
    }
    
    def __init__(self, docx_dir: str):
        self.docx_dir = Path(docx_dir)
        self.causes_by_code: Dict[str, Dict] = {}  # code -> {tr: {...}, en: {...}}
        
    def detect_language(self, text: str) -> Optional[str]:
        """Metinden dili tespit eder."""
        for lang, markers in self.LANGUAGE_MARKERS.items():
            if any(marker in text for marker in markers):
                return lang
        return None
    
    def parse_docx(self, docx_path: Path, language: str) -> Dict[str, Dict]:
        """Tek bir DOCX dosyasını ayrıştırır."""
        print(f"\n📄 {language.upper()} dosyası yükleniyor: {docx_path.name}")
        document = Document(docx_path)
        print(f"✓ Toplam {len(document.paragraphs)} paragraf yüklendi")
        
        causes = {}
        current_cause = None
        current_section = None
        waiting_for_definition = False
        
        for para in document.paragraphs:
            text = para.text.strip()
            # Non-breaking space ve diğer whitespace karakterlerini normalize et
            text = text.replace('\xa0', ' ').replace('\n', ' ')
            
            if not text:
                continue
            
            # Kod pattern'i: A1.1, B2.3, vb.
            code_match = re.match(r'^([A-Z]\d+\.\d+)\s+(.+)$', text)
            
            if code_match:
                # Yeni bir cause başlıyor - öncekini kaydet
                if current_cause:
                    causes[current_cause['code']] = current_cause
                
                code = code_match.group(1)
                title = code_match.group(2).strip()
                
                current_cause = {
                    'code': code,
                    'title': title,
                    'definition': '',
                    'selection_criteria': None,
                    'typical_examples': [],
                    'exclusion_conditions': []
                }
                current_section = None
                waiting_for_definition = True
                
            elif current_cause:
                # "Seç eğer:" / "Choose if:" ile başlayan kriter
                if text.startswith('Seç eğer:') or text.startswith('Choose if:') or 'Seç eğer:' in text or 'Choose if:' in text:
                    current_section = 'selection_criteria'
                    waiting_for_definition = False
                    
                    # Kriter metnini çıkar
                    if 'Seç eğer:' in text:
                        criteria_text = text.split('Seç eğer:')[1].strip()
                    elif 'Choose if:' in text:
                        criteria_text = text.split('Choose if:')[1].strip()
                    else:
                        criteria_text = text
                    
                    # "→ Tipik:" / "→ Typical:" varsa onu ayır
                    if '→ Tipik:' in criteria_text or '→ Typical:' in criteria_text:
                        if '→ Tipik:' in criteria_text:
                            parts = criteria_text.split('→ Tipik:')
                        else:
                            parts = criteria_text.split('→ Typical:')
                        
                        current_cause['selection_criteria'] = parts[0].strip()
                        if len(parts) > 1:
                            examples_text = parts[1].strip()
                            if examples_text:
                                current_cause['typical_examples'].append(examples_text)
                    else:
                        current_cause['selection_criteria'] = criteria_text
                    continue
                
                # İlk paragraf tanımdır
                if waiting_for_definition and not text.startswith('→'):
                    current_cause['definition'] = text
                    waiting_for_definition = False
                    continue
                
                # "→ Tipik:" / "→ Typical:" ile başlayan örnekler
                if text.startswith('→ Tipik:') or text.startswith('→ Typical:'):
                    current_section = 'examples'
                    if '→ Tipik:' in text:
                        examples_text = text.replace('→ Tipik:', '').strip()
                    else:
                        examples_text = text.replace('→ Typical:', '').strip()
                    if examples_text:
                        current_cause['typical_examples'].append(examples_text)
                    continue
                
                # "Bu değil eğer:" / "Not this if:" ile başlayan exclusion bölümü
                if text.startswith('Bu değil eğer:') or text.startswith('Not this if:'):
                    current_section = 'exclusion'
                    continue
                
                # Exclusion bölümündeki satırlar
                if current_section == 'exclusion':
                    if '→' in text:
                        parts = text.split('→')
                        condition = parts[0].strip()
                        redirect_info = parts[1].strip() if len(parts) > 1 else ''
                        
                        # Redirect code'u bul
                        redirect_match = re.search(r'([A-Z]\d+\.\d+)', redirect_info)
                        redirect_code = redirect_match.group(1) if redirect_match else ''
                        
                        if condition and redirect_code:
                            exclusion = {
                                'condition': condition,
                                'redirect_code': redirect_code,
                                'reason': redirect_info,
                                'language': language
                            }
                            current_cause['exclusion_conditions'].append(exclusion)
        
        # Son cause'u ekle
        if current_cause:
            causes[current_cause['code']] = current_cause
        
        print(f"✓ {len(causes)} cause ayrıştırıldı")
        return causes
    
    def merge_languages(self) -> Taxonomy:
        """Tüm dillerdeki cause'ları birleştirir."""
        print("\n🔄 Diller birleştiriliyor...")
        
        cause_objects = []
        
        for code in sorted(self.causes_by_code.keys()):
            lang_data = self.causes_by_code[code]
            
            # Cause type'ı belirle
            cause_type = self._determine_cause_type(code)
            
            # Her dil için LocalizedContent oluştur
            content = {}
            all_exclusions = []
            
            for lang, data in lang_data.items():
                localized = LocalizedContent(
                    title=data['title'],
                    definition=data['definition'],
                    selection_criteria=data.get('selection_criteria'),
                    typical_examples=data.get('typical_examples', [])
                )
                content[lang] = localized
                
                # Exclusion'ları topla
                for exc in data.get('exclusion_conditions', []):
                    all_exclusions.append(ExclusionCondition(**exc))
            
            # Cause objesi oluştur
            cause = Cause(
                code=code,
                cause_type=cause_type,
                content=content,
                exclusion_conditions=all_exclusions,
                related_codes=[],
                keywords={},
                severity_indicators=[],
                industry_contexts=[]
            )
            
            cause_objects.append(cause)
        
        print(f"✓ {len(cause_objects)} cause birleştirildi")
        
        taxonomy = Taxonomy(causes=cause_objects)
        return taxonomy
    
    def _determine_cause_type(self, code: str) -> str:
        """Kod prefix'ine göre cause type'ı belirler."""
        prefix = code[0]
        
        type_mapping = {
            'A': 'immediate_cause',
            'B': 'root_cause',
            'C': 'root_cause',
            'D': 'root_cause',
            'E': 'root_cause'
        }
        
        return type_mapping.get(prefix, 'unknown')
    
    def parse_all(self) -> Taxonomy:
        """Tüm DOCX dosyalarını ayrıştırır."""
        print("=" * 70)
        print("🚀 Multilingual Taxonomy Parser")
        print("=" * 70)
        
        # DOCX dosyalarını bul
        docx_files = list(self.docx_dir.glob('*.docx'))
        print(f"\n📁 {len(docx_files)} DOCX dosyası bulundu:")
        for f in docx_files:
            print(f"   - {f.name}")
        
        # Her dosyayı ayrıştır
        for docx_file in docx_files:
            # Dosya adından dili tespit et
            if 'Türkçe' in docx_file.name or 'turkce' in docx_file.name.lower():
                language = 'tr'
            elif 'İngilizce' in docx_file.name or 'ingilizce' in docx_file.name.lower() or 'English' in docx_file.name:
                language = 'en'
            else:
                # İçerikten tespit et
                doc = Document(docx_file)
                sample_text = ' '.join([p.text for p in doc.paragraphs[:10]])
                language = self.detect_language(sample_text) or 'unknown'
            
            if language == 'unknown':
                print(f"⚠️  {docx_file.name} için dil tespit edilemedi, atlanıyor...")
                continue
            
            # Dosyayı ayrıştır
            causes = self.parse_docx(docx_file, language)
            
            # Mevcut cause'larla birleştir
            for code, data in causes.items():
                if code not in self.causes_by_code:
                    self.causes_by_code[code] = {}
                self.causes_by_code[code][language] = data
        
        # Tüm dilleri birleştir
        taxonomy = self.merge_languages()
        
        return taxonomy
    
    def save_to_json(self, taxonomy: Taxonomy, output_path: str):
        """Taxonomy objesini JSON dosyasına kaydeder."""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Pydantic model_dump kullanarak JSON'a dönüştür
        taxonomy_dict = taxonomy.model_dump(by_alias=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(taxonomy_dict, f, ensure_ascii=False, indent=2)
        
        print(f"\n✅ Çok dilli taxonomy JSON'a kaydedildi: {output_path}")
        print(f"   📊 Toplam cause sayısı: {len(taxonomy.causes)}")
        
        # İstatistikler
        immediate_causes = [c for c in taxonomy.causes if c.cause_type == 'immediate_cause']
        root_causes = [c for c in taxonomy.causes if c.cause_type == 'root_cause']
        
        # Dil istatistikleri
        lang_stats = {}
        for cause in taxonomy.causes:
            for lang in cause.content.keys():
                lang_stats[lang] = lang_stats.get(lang, 0) + 1
        
        print(f"   🔵 Immediate causes: {len(immediate_causes)}")
        print(f"   🔴 Root causes: {len(root_causes)}")
        print(f"   🌍 Diller:")
        for lang, count in sorted(lang_stats.items()):
            print(f"      - {lang.upper()}: {count} cause")


def main():
    """Ana çalıştırma fonksiyonu."""
    # Dosya yolları
    project_root = Path(__file__).parent.parent.parent
    docx_dir = project_root / "knowlodge_base"
    output_path = project_root / "rag_pipeline" / "data" / "processed" / "taxonomy_multilingual.json"
    
    try:
        # Parser'ı oluştur ve çalıştır
        parser = MultilingualTaxonomyParser(str(docx_dir))
        taxonomy = parser.parse_all()
        
        # JSON'a kaydet
        parser.save_to_json(taxonomy, str(output_path))
        
        print("\n" + "=" * 70)
        print("✅ İşlem başarıyla tamamlandı!")
        print("=" * 70)
        
    except Exception as e:
        print(f"\n❌ Hata oluştu: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
