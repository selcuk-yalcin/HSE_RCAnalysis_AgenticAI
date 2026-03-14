"""
Taxonomy Parser - DOCX'ten Yapılandırılmış JSON'a Dönüştürücü
=============================================================

Bu script, knowlodge_base/Türkçe_Taksonomi.docx dosyasını okur ve
rag_pipeline/schemas/cause_models.py'deki Pydantic modellerine göre
yapılandırılmış bir JSON dosyası oluşturur.

Kullanım:
    python rag_pipeline/parsing/parse_taxonomy.py
"""

import sys
import json
import re
from pathlib import Path
from typing import List, Dict, Optional
from docx import Document

# Proje kök dizinini Python path'e ekle
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

from rag_pipeline.schemas.cause_models import Cause, ExclusionCondition, Taxonomy


class TaxonomyParser:
    """DOCX dosyasından taksonomi verilerini ayrıştırır."""
    
    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        self.document = None
        self.causes: List[Cause] = []
        
    def load_document(self):
        """DOCX dosyasını yükler."""
        if not self.docx_path.exists():
            raise FileNotFoundError(f"DOCX dosyası bulunamadı: {self.docx_path}")
        
        print(f"📄 DOCX dosyası yükleniyor: {self.docx_path}")
        self.document = Document(self.docx_path)
        print(f"✓ Toplam {len(self.document.paragraphs)} paragraf yüklendi")
        
    def parse(self) -> Taxonomy:
        """DOCX içeriğini ayrıştırır ve Taxonomy objesi döndürür."""
        self.load_document()
        
        current_cause = None
        current_section = None
        waiting_for_definition = False
        
        for para in self.document.paragraphs:
            text = para.text.strip()
            # Non-breaking space ve diğer whitespace karakterlerini normalize et
            text = text.replace('\xa0', ' ').replace('\n', ' ')
            
            if not text:
                continue
            
            # Kod pattern'i: A1.1, B2.3, vb.
            code_match = re.match(r'^([A-Z]\d+\.\d+)\s+(.+)$', text)
            
            if code_match:
                # Yeni bir cause başlıyor - öncekini kaydet
                if current_cause:
                    self.causes.append(current_cause)
                
                code = code_match.group(1)
                title = code_match.group(2).strip()
                
                # Cause type'ı koddan belirle
                cause_type = self._determine_cause_type(code)
                
                current_cause = {
                    'code': code,
                    'title': title,
                    'definition': '',
                    'cause_type': cause_type,
                    'selection_criteria': None,
                    'typical_examples': [],
                    'exclusion_conditions': [],
                    'keywords': [],
                    'related_codes': [],
                    'severity_indicators': [],
                    'industry_contexts': []
                }
                current_section = None
                waiting_for_definition = True
                
            elif current_cause:
                # Mevcut cause'a ait içerik
                
                # "Seç eğer:" ile başlayan kriter (önce bu kontrol edilmeli)
                if text.startswith('Seç eğer:') or 'Seç eğer:' in text:
                    current_section = 'selection_criteria'
                    waiting_for_definition = False  # Artık definition beklemeyi bırak
                    
                    # "Seç eğer:" kısmını çıkar
                    if 'Seç eğer:' in text:
                        criteria_text = text.split('Seç eğer:')[1].strip()
                    else:
                        criteria_text = text.replace('Seç eğer:', '').strip()
                    
                    # "→ Tipik:" varsa onu ayır
                    if '→ Tipik:' in criteria_text:
                        parts = criteria_text.split('→ Tipik:')
                        current_cause['selection_criteria'] = parts[0].strip()
                        # Örnekleri de ekle
                        if len(parts) > 1:
                            examples_text = parts[1].strip()
                            if examples_text:
                                current_cause['typical_examples'].append(examples_text)
                    else:
                        current_cause['selection_criteria'] = criteria_text
                    continue
                
                # İlk paragraf tanımdır (Seç eğer'den sonra kontrol et)
                if waiting_for_definition and not text.startswith('→'):
                    current_cause['definition'] = text
                    waiting_for_definition = False
                    continue
                
                # "→ Tipik:" ile başlayan örnekler (ayrı satırda)
                if text.startswith('→ Tipik:'):
                    current_section = 'examples'
                    examples_text = text.replace('→ Tipik:', '').strip()
                    if examples_text:
                        current_cause['typical_examples'].append(examples_text)
                    continue
                
                # "Bu değil eğer:" ile başlayan exclusion bölümü
                if text.startswith('Bu değil eğer:'):
                    current_section = 'exclusion'
                    continue
                
                # Exclusion bölümündeki satırlar
                if current_section == 'exclusion':
                    # "→" ile başlayan redirect satırları
                    if '→' in text:
                        parts = text.split('→')
                        condition = parts[0].strip()
                        redirect_info = parts[1].strip() if len(parts) > 1 else ''
                        
                        # Redirect code'u bul - önce parantez içinde, yoksa başta
                        redirect_match = re.search(r'([A-Z]\d+\.\d+)', redirect_info)
                        redirect_code = redirect_match.group(1) if redirect_match else ''
                        
                        if condition and redirect_code:
                            exclusion = ExclusionCondition(
                                condition=condition,
                                redirect_code=redirect_code,
                                reason=redirect_info
                            )
                            current_cause['exclusion_conditions'].append(exclusion)
                    # Yeni bir cause başlangıcı değilse devam et
                    elif not re.match(r'^[A-Z]\d+\.\d+\s', text):
                        # Çok satırlı exclusion condition olabilir
                        pass
        
        # Son cause'u ekle
        if current_cause:
            self.causes.append(current_cause)
        
        print(f"\n✓ Toplam {len(self.causes)} cause ayrıştırıldı")
        
        # Pydantic modellerine dönüştür
        cause_objects = []
        for cause_dict in self.causes:
            try:
                cause_obj = Cause(**cause_dict)
                cause_objects.append(cause_obj)
            except Exception as e:
                print(f"⚠️  {cause_dict['code']} dönüştürülürken hata: {e}")
        
        taxonomy = Taxonomy(
            version="1.0",
            last_updated=None,  # Otomatik olarak şimdi atanacak
            causes=cause_objects
        )
        
        return taxonomy
    
    def _determine_cause_type(self, code: str) -> str:
        """Kod prefix'ine göre cause type'ı belirler."""
        prefix = code[0]
        
        type_mapping = {
            'A': 'immediate_cause',
            'B': 'root_cause',
            'C': 'root_cause',
            'D': 'root_cause',
            'E': 'root_cause'
        }
        
        return type_mapping.get(prefix, 'unknown')
    
    def save_to_json(self, taxonomy: Taxonomy, output_path: str):
        """Taxonomy objesini JSON dosyasına kaydeder."""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Pydantic model_dump kullanarak JSON'a dönüştür
        taxonomy_dict = taxonomy.model_dump(by_alias=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(taxonomy_dict, f, ensure_ascii=False, indent=2)
        
        print(f"\n✅ Taxonomy JSON'a kaydedildi: {output_path}")
        print(f"   📊 Toplam cause sayısı: {len(taxonomy.causes)}")
        
        # İstatistikler
        immediate_causes = [c for c in taxonomy.causes if c.cause_type == 'immediate_cause']
        root_causes = [c for c in taxonomy.causes if c.cause_type == 'root_cause']
        
        print(f"   🔵 Immediate causes: {len(immediate_causes)}")
        print(f"   🔴 Root causes: {len(root_causes)}")


def main():
    """Ana çalıştırma fonksiyonu."""
    print("=" * 70)
    print("🚀 Taxonomy Parser - DOCX'ten JSON'a Dönüştürücü")
    print("=" * 70)
    
    # Dosya yolları
    project_root = Path(__file__).parent.parent.parent
    docx_path = project_root / "knowlodge_base" / "Türkçe_Taksonomi.docx"
    output_path = project_root / "rag_pipeline" / "data" / "processed" / "taxonomy.json"
    
    try:
        # Parser'ı oluştur ve çalıştır
        parser = TaxonomyParser(str(docx_path))
        taxonomy = parser.parse()
        
        # JSON'a kaydet
        parser.save_to_json(taxonomy, str(output_path))
        
        print("\n" + "=" * 70)
        print("✅ İşlem başarıyla tamamlandı!")
        print("=" * 70)
        
    except Exception as e:
        print(f"\n❌ Hata oluştu: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
