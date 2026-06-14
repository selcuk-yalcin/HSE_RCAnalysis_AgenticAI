"""
SkillBasedDocxAgent V2 - OpenRouter LLM + python-docx ile Profesyonel HSE Raporu
========================================================================================

MİMARİ:
  RootCauseAgentV2 → JSON → OpenRouter LLM (içerik üretir) → python-docx (DOCX oluşturur)

AVANTAJLAR:
  - OpenRouter üzerinden seçili modeli kullanır
  - python-docx ile kesin, güvenilir DOCX oluşturma
  - 10-20 sayfalık profesyonel rapor
  - HSE renk şeması: koyu mavi, kırmızı, turuncu, yeşil kutular/tablolar

GEREKSİNİMLER:
  pip install requests python-docx

ORTAM DEĞİŞKENLERİ:
  OPENROUTER_API_KEY = "sk-or-v1-..."
"""

from __future__ import annotations

import requests
import json
import os
import sys
import re
import time
import copy
from typing import Dict, List, Optional, Any
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv

from .model_constants import resolve_openrouter_docx_model

load_dotenv()

try:
    from agents.json_parser import extract_json_from_response
except ImportError:
    from json_parser import extract_json_from_response

try:
    from agents.report_text_sanitize import (
        format_report_html_rich,
        sanitize_report_text,
        set_report_text_policy,
        short_incident_summary,
        strip_root_cause_label_prefix,
        taxonomy_display_title,
    )
    from shared.report_layout_config import resolve_report_layout
    from shared.report_i18n import (
        apply_shell_language,
        get_report_lang,
        report_label,
        set_report_lang,
        shell_fallback_replacements,
    )
except ImportError:
    from .report_text_sanitize import (
        format_report_html_rich,
        sanitize_report_text,
        set_report_text_policy,
        short_incident_summary,
        strip_root_cause_label_prefix,
        taxonomy_display_title,
    )
    from shared.report_layout_config import resolve_report_layout
    from shared.report_i18n import (
        apply_shell_language,
        get_report_lang,
        report_label,
        set_report_lang,
        shell_fallback_replacements,
    )

# Geriye dönük importlar
strip_hse_codes = sanitize_report_text


def _direct_cause_sentence(text: str, lang_code: str = "tr") -> str:
    """Doğrudan neden ifadesini düşük (yarım) cümle olarak basmak yerine tam cümleye çevirir."""
    t = re.sub(r"\s+", " ", str(text or "").strip())
    if not t:
        return ""
    if t[-1] in ".!?":
        return t
    t = t.rstrip(" ,;:")
    if (lang_code or "tr").lower().startswith("en"):
        return f"The direct cause of this branch was identified as: {t}."
    return f"Bu dalın doğrudan nedeni, {t} olarak belirlenmiştir."


def _default_barsel_code_system(lang_code: str = "tr") -> List[Dict[str, str]]:
    """BARSEL/HSG245 3.2 kod tablosu: A=Davranış, B=Koşullar, C=Kişisel, D=Organizasyonel."""
    if (lang_code or "tr").lower().startswith("en"):
        return [
            {
                "code": "A",
                "category": "Behavior",
                "description": "Observable actions: procedure violations, PPE use, unsafe acts",
            },
            {
                "code": "B",
                "category": "Conditions",
                "description": "Equipment, environment, energy sources, and physical workplace conditions",
            },
            {
                "code": "C",
                "category": "Personal",
                "description": "Individual capacity, competence, skill, and performance factors",
            },
            {
                "code": "D",
                "category": "Organizational",
                "description": "Management systems, procedures, training, supervision, and oversight gaps",
            },
        ]
    return [
        {
            "code": "A",
            "category": "Davranış",
            "description": "Gözlemlenebilir eylemler: prosedür ihlali, KKD kullanımı, güvensiz davranış",
        },
        {
            "code": "B",
            "category": "Koşullar",
            "description": "Ekipman, ortam, enerji kaynağı ve fiziksel iş yeri koşulları",
        },
        {
            "code": "C",
            "category": "Kişisel",
            "description": "Bireysel kapasite, yeterlilik, beceri ve performans faktörleri",
        },
        {
            "code": "D",
            "category": "Organizasyonel",
            "description": "Yönetim sistemleri, prosedür, eğitim, gözetim ve denetim eksiklikleri",
        },
    ]


def _normalize_analysis_method(method: Optional[Dict]) -> Dict:
    """3.2 kod tablosunu resmi A/B/C/D etiketleriyle hizala."""
    m = dict(method or {})
    lang = get_report_lang()
    canonical = {row["code"]: row for row in _default_barsel_code_system(lang)}
    rows = m.get("code_system") or []
    if not rows:
        m["code_system"] = list(canonical.values())
        return m
    fixed: List[Dict[str, str]] = []
    seen: set[str] = set()
    for row in rows:
        letter = str(row.get("code") or "").strip().upper()[:1]
        if letter not in canonical or letter in seen:
            continue
        seen.add(letter)
        base = canonical[letter]
        fixed.append(
            {
                "code": letter,
                "category": base["category"],
                "description": str(row.get("description") or base["description"]).strip(),
            }
        )
    for letter in ("A", "B", "C", "D"):
        if letter not in seen:
            fixed.append(canonical[letter])
    m["code_system"] = sorted(fixed, key=lambda r: r.get("code", ""))
    return m


def _resolve_openrouter_chat_completions_url() -> str:
    """OPENROUTER_BASE_URL veya tam URL; çift /v1 ve yanlış path hatalarını azaltır."""
    explicit = (os.getenv("OPENROUTER_CHAT_COMPLETIONS_URL") or "").strip()
    if explicit:
        return explicit.rstrip("/")
    base = (os.getenv("OPENROUTER_BASE_URL") or "https://openrouter.ai/api/v1").strip().rstrip("/")
    while "/v1/v1" in base:
        base = base.replace("/v1/v1", "/v1", 1)
    if base.endswith("/chat/completions"):
        return base
    return f"{base}/chat/completions"


# python-docx imports
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# 
# RENK PALETİ
# 
COLOR = {
    "dark_blue":  RGBColor(0x1B, 0x3A, 0x5C),
    "mid_blue":   RGBColor(0x2E, 0x6D, 0xA4),
    "light_blue": RGBColor(0xD6, 0xE4, 0xF0),
    "red":        RGBColor(0xC0, 0x39, 0x2B),
    "orange":     RGBColor(0xE6, 0x7E, 0x22),
    "green":      RGBColor(0x27, 0xAE, 0x60),
    "light_grey": RGBColor(0xF5, 0xF5, 0xF5),
    "white":      RGBColor(0xFF, 0xFF, 0xFF),
    "black":      RGBColor(0x00, 0x00, 0x00),
    "dark_grey":  RGBColor(0x44, 0x44, 0x44),
}

ROOT_CAUSE_COLORS = [
    COLOR["red"],
    COLOR["orange"],
    COLOR["green"],
    COLOR["mid_blue"],
]


def _label(lang_code: str, key: str) -> str:
    return report_label(lang_code, key)


def _L(key: str, lang_code: Optional[str] = None) -> str:
    return report_label(lang_code or get_report_lang(), key)


def _translate_html_static_labels(html: str, lang_code: str) -> str:
    return apply_shell_language(html, lang_code)


def _translate_docx_static_labels(doc: Document, lang_code: str) -> None:
    if (lang_code or "tr").lower() == "tr":
        return
    pairs = shell_fallback_replacements(lang_code)

    def _translate_text(text: str) -> str:
        if not text:
            return text
        out = text
        for tr_text, en_text in pairs.items():
            out = out.replace(tr_text, en_text)
        return out

    for paragraph in doc.paragraphs:
        original = paragraph.text
        translated = _translate_text(original)
        if translated != original:
            paragraph.text = translated

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original = cell.text
                translated = _translate_text(original)
                if translated != original:
                    cell.text = translated


# 
# CLAUDE CONTENT PROMPT
# 
CONTENT_SYSTEM_PROMPT = """Sen bir HSE (İş Sağlığı ve Güvenliği) uzmanısın.
Sana bir kök neden analizi ham verisi gelecek.
Bu veriyi kullanarak raporun TÜM İÇERİĞİNİ üreteceksin.

Sadece JSON formatında çıktı ver. Başka hiçbir şey yazma.

Üretmen gereken JSON yapısı:
{
  "cover": {
    "title": "KÖK NEDEN ANALİZİ RAPORU",
    "subtitle": "Profesyonel Araştırma ve Analiz Raporu",
    "ref_no": "...",
    "date": "...",
    "location": "...",
    "incident_type": "...",
    "confidentiality": "GİZLİ - SADECE YETKİLİ PERSONELİN ERİŞİMİNE AÇIKTIR",
    "incident_summary_short": "Olayın 2-3 cümle özeti - tekrarlardan kaçın, öz ve net yaz"
  },
  "executive_summary": {
    "what_happened": "Ne oldu - 1-2 paragraf; yer, ekipman ve etkilenen kişi bilgisini burada doğal cümlelerle ver (ayrı satırda tekrarlama yok)",
    "where_happened": "",
    "who_affected": "",
    "immediate_response": "İlk müdahale - 1 paragraf",
    "key_findings": [
      "Bulgu 1 - kısa",
      "Bulgu 2 - kısa",
      "Bulgu 3 - kısa"
    ],
    "immediate_actions": [
      {"action": "Acil eylem 1", "responsible": "Sorumlu", "status": "Tamamlandı"},
      {"action": "Acil eylem 2", "responsible": "Sorumlu", "status": "Devam ediyor"},
      {"action": "Acil eylem 3", "responsible": "Sorumlu", "status": "Planlandı"},
      {"action": "Acil eylem 4", "responsible": "Sorumlu", "status": "Tamamlandı"}
    ]
  },
  "incident_details": {
    "info_table": {
      "Olay Referans No": "...",
      "Tarih": "...",
      "Saat": "...",
      "Lokasyon": "...",
      "Bölüm/Hat": "...",
      "Operatör/Çalışan": "...",
      "Vardiya": "...",
      "Ekipman": "...",
      "Malzeme/Madde": "...",
      "Hava Koşulları": "...",
      "Aydınlatma": "...",
      "Kişisel Koruyucu Ekipman": "..."
    },
    "event_table": {
      "Olay Tipi": "...",
      "Yaralanma/Hasar Durumu": "...",
      "Etkilenen Kişi Sayısı": "...",
      "Hasar Seviyesi": "...",
      "İlk Tanık": "...",
      "Acil Servis Çağrıldı mı": "...",
      "Yatış/Taburculuk": "..."
    },
    "timeline": [
      {"time": "00:00", "event": "Olaydan önce durum açıklaması"},
      {"time": "00:05", "event": "Olayın başlangıcı"},
      {"time": "00:10", "event": "Olay anı"},
      {"time": "00:15", "event": "İlk müdahale"},
      {"time": "00:30", "event": "Acil servis/yönetim bildirim"},
      {"time": "01:00", "event": "Durum kontrolü ve raporlama"}
    ],
    "severity": {
      "actual_harm": "...",
      "potential_harm": "...",
      "investigation_level": "..."
    }
  },
  "analysis_method": {
    "methodology_description": "Kök neden analizi yöntemi ve bu olayda nasıl uygulandı - 2 paragraf",
    "five_why_explanation": "5-Why tekniği nasıl uygulandı - 2 paragraf",
    "code_system": [
      {"code": "A", "category": "Davranış", "description": "Gözlemlenebilir eylemler: prosedür ihlali, KKD kullanımı, güvensiz davranış"},
      {"code": "B", "category": "Koşullar", "description": "Ekipman, ortam, enerji kaynağı ve fiziksel iş yeri koşulları"},
      {"code": "C", "category": "Kişisel", "description": "Bireysel kapasite, yeterlilik, beceri ve performans faktörleri"},
      {"code": "D", "category": "Organizasyonel", "description": "Yönetim sistemleri, prosedür, eğitim, gözetim ve denetim eksiklikleri"}
    ],
    "team_members": [
      {"name": "HSE Uzmanı", "role": "Baş Araştırmacı", "date": "..."},
      {"name": "Üretim Müdürü", "role": "Departman Temsilcisi", "date": "..."},
      {"name": "Bakım Mühendisi", "role": "Teknik Uzman", "date": "..."},
      {"name": "İK Yöneticisi", "role": "İnsan Kaynakları Temsilcisi", "date": "..."},
      {"name": "Vardiya Amiri", "role": "Operasyonel Tanık", "date": "..."}
    ]
  },
  "branches": [
    {
      "branch_number": 1,
      "branch_title": "KRİTİK FAKTÖR 1 - Mühendislik / Tasarım ve Teknik Sistemler",
      "initial_condition": "Bu faktörün başlangıç koşulu - 1 paragraf",
      "direct_cause": "Doğrudan nedenin kısa açıklaması",
      "why_chain": [
        {"number": 1, "question": "(sistem doldurur — boş bırak)", "answer": "", "code": "", "category": ""},
        {"number": 2, "question": "", "answer": "", "code": "", "category": ""},
        {"number": 3, "question": "", "answer": "", "code": "", "category": ""},
        {"number": 4, "question": "", "answer": "", "code": "", "category": ""},
        {"number": 5, "question": "", "answer": "", "code": "", "category": ""}
      ],
      "root_cause_title": "Kök Neden 1 başlığı",
      "root_cause_detail": "Kök nedenin çok detaylı açıklaması - 3-4 cümle",
      "organizational_factors": [
        "Organizasyonel faktör 1 - detaylı",
        "Organizasyonel faktör 2 - detaylı",
        "Organizasyonel faktör 3 - detaylı",
        "Organizasyonel faktör 4 - detaylı"
      ]
    }
  ],
  "root_causes": [
    {
      "number": 1,
      "title": "Kök Neden Başlığı",
      "category": "Organizasyonel",
      "detailed_description": "3-4 paragraf çok detaylı açıklama",
      "impacts": ["Etki 1", "Etki 2", "Etki 3", "Etki 4"],
      "contributing_organizations": "Hangi organizasyonel birimler bu nedenle ilişkili"
    }
  ],
  "meta_root_cause": {
    "exists": true,
    "code": "D8.2",
    "title": "Meta Kök Neden Başlığı (Tüm Dalların Ortak Paydası)",
    "description": "3-4 paragraf: Tüm dalların ortak paydası olan üst-seviye organizasyonel zayıflık",
    "synthesized_from": ["D2.1", "D6.3", "D4.10"],
    "systemic_weakness": "Hangi yönetim sistemi zayıflığı tüm nedenleri doğurdu",
    "strategic_implications": [
      "Stratejik sonuç 1",
      "Stratejik sonuç 2",
      "Stratejik sonuç 3"
    ]
  },
  "contributing_factors": [
    {"factor_type": "İletişim Eksikliği", "description": "Güvenlik prosedürlerinin önemi ve zorunluluğu çalışanlara etkin şekilde iletilmemiş", "impact_level": "Yüksek"},
    {"factor_type": "Eğitim Yetersizliği", "description": "LOTO prosedürü ve elektrik güvenliği eğitimleri yetersiz veya etkin değil", "impact_level": "Yüksek"},
    {"factor_type": "Gözetim Eksikliği", "description": "Saha gözetim ve denetim mekanizması tanımsız ve güvensiz", "impact_level": "Yüksek"},
    {"factor_type": "Kültürel Faktörler", "description": "Üretim öncelikli kültür, güvenlik sapmalarına tolere edilmesi", "impact_level": "Yüksek"}
  ],
  "corrective_actions": [
    {"no": 1, "action": "LOTO prosedürü tüm elektrik panolarında zorunlu hale getirilmesi", "priority": "ACİL", "responsible": "HSE Yöneticisi", "deadline": "1 hafta", "kpi": "LOTO uygulanma oranı %100"},
    {"no": 2, "action": "Enerjili çalışmalarda KKD kullanımı denetim sistemi kurulması", "priority": "ACİL", "responsible": "Üretim Müdürü", "deadline": "2 hafta", "kpi": "KKD kullanım uyumu %100"},
    {"no": 3, "action": "Elektrik güvenliği ve LOTO zorunlu eğitimi", "priority": "YÜKSEK", "responsible": "İK Müdürü", "deadline": "1 ay", "kpi": "Eğitim tamamlama %100"},
    {"no": 4, "action": "Saha gözetim ve denetim programı oluşturulması", "priority": "YÜKSEK", "responsible": "Vardiya Amiri", "deadline": "1 ay", "kpi": "Denetim frekansı günlük"},
    {"no": 5, "action": "İş durdurma yetkisi mekanizması kurulması", "priority": "YÜKSEK", "responsible": "Tesis Müdürü", "deadline": "2 ay", "kpi": "Yetki kullanım oranı %100"},
    {"no": 6, "action": "Üretim-güvenlik dengesi politikası oluşturulması", "priority": "ORTA", "responsible": "Genel Müdür", "deadline": "2 ay", "kpi": "Politika uygulama %100"},
    {"no": 7, "action": "Güvenlik sapmaları raporlama sistemi kurulması", "priority": "ORTA", "responsible": "HSE Uzmanı", "deadline": "3 ay", "kpi": "Sapma raporlama oranı"},
    {"no": 8, "action": "LOTO ekipman temini ve kontrolü", "priority": "ORTA", "responsible": "Bakım Müdürü", "deadline": "3 ay", "kpi": "Ekipman yeterliliği %100"},
    {"no": 9, "action": "Güvenlik kültürü geliştirme programı", "priority": "ORTA", "responsible": "İK Müdürü", "deadline": "3 ay", "kpi": "Kültür anket skoru artışı"},
    {"no": 10, "action": "Liderlik taahhüdü güçlendirme", "priority": "DÜŞÜK", "responsible": "Üst Yönetim", "deadline": "6 ay", "kpi": "Liderlik görünürlüğü"},
    {"no": 11, "action": "Elektrik güvenlik standartları revizyonu", "priority": "DÜŞÜK", "responsible": "Teknik Direktör", "deadline": "6 ay", "kpi": "Standart güncellik"},
    {"no": 12, "action": "Güvenlik performans takip sistemi", "priority": "DÜŞÜK", "responsible": "HSE Yöneticisi", "deadline": "12 ay", "kpi": "Sistem etkinliği"}
  ],
  "lessons_learned": {
    "what_to_do": [
      "Ders 1 - Ne yapılmalı: detaylı açıklama",
      "Ders 2 - Ne yapılmalı: detaylı açıklama",
      "Ders 3 - Ne yapılmalı: detaylı açıklama",
      "Ders 4 - Ne yapılmalı: detaylı açıklama"
    ],
    "long_term": [
      "Uzun vadeli çözüm 1: detaylı açıklama",
      "Uzun vadeli çözüm 2: detaylı açıklama",
      "Uzun vadeli çözüm 3: detaylı açıklama",
      "Uzun vadeli çözüm 4: detaylı açıklama"
    ],
    "communication": [
      "İletişim planı 1: detaylı açıklama",
      "İletişim planı 2: detaylı açıklama",
      "İletişim planı 3: detaylı açıklama"
    ],
    "training": [
      "Eğitim programı 1: detaylı açıklama",
      "Eğitim programı 2: detaylı açıklama",
      "Eğitim programı 3: detaylı açıklama",
      "Eğitim programı 4: detaylı açıklama"
    ]
  },
  "conclusion": {
    "overall_assessment": "Genel değerlendirme - 3-4 paragraf kapsamlı",
    "short_term_measures": [
      "Kısa vade önlem 1 (1-2 ay): detaylı",
      "Kısa vade önlem 2 (1-2 ay): detaylı",
      "Kısa vade önlem 3 (1-2 ay): detaylı",
      "Kısa vade önlem 4 (1-2 ay): detaylı"
    ],
    "long_term_improvements": [
      "Uzun vade iyileştirme 1 (3-12 ay): detaylı",
      "Uzun vade iyileştirme 2 (3-12 ay): detaylı",
      "Uzun vade iyileştirme 3 (3-12 ay): detaylı",
      "Uzun vade iyileştirme 4 (3-12 ay): detaylı"
    ],
    "comparison_table": [
      {"criterion": "Risk Seviyesi", "current": "Yüksek", "target": "Düşük"},
      {"criterion": "Prosedür Uyum Oranı", "current": "%60", "target": "%95"},
      {"criterion": "Eğitim Kapsamı", "current": "Temel", "target": "Kapsamlı"},
      {"criterion": "Bakım Periyodu", "current": "Reaktif", "target": "Proaktif"},
      {"criterion": "Denetim Sıklığı", "current": "Aylık", "target": "Haftalık"}
    ]
  }
}

KURALLAR:
- DİL KURALI: Ham verinin yazıldığı dili tespit et. Arapça ise ARAPÇA, İngilizce ise İNGİLİZCE, İspanyolca ise İSPANYOLCA, Fransızca ise FRANSIZCA, Türkçe ise TÜRKÇE yaz. Dilleri ASLA karıştırma. JSON key'leri İngilizce kalır, tüm VALUE'lar ham verinin dilinde olur.
- Her alan ham veriden türetilmeli
- Kısa cevaplar değil, DETAYLI açıklamalar
- branches dizisi ham verideki tüm dalları içermeli
- root_causes dizisi ham verideki tüm kök nedenleri içermeli
- Sınıflandırma / HSG kodları (ör. D4.1, C3.1, H-1.5, K-01), parantez içi kodlar veya "Birincil Kod:" gibi etiketler narrative metinlerde, kök neden açıklamalarında ve why_chain soru-cevaplarında YAZILMAYACAK. JSON şemasındaki code alanlarını boş string bırak veya kullanma.
- executive_summary: where_happened ve who_affected alanlarını her zaman boş string "" bırak; yer ve kişi bilgisini yalnızca what_happened içinde anlat.
- Olay özeti yalnızca cover.incident_summary_short içinde olmalı (2-3 cümle). incident_details.event_table içine "Özet" veya uzun anlatım EKLEME.
- why_chain: BU ALANI YENİDEN YAZMA — boş bırak veya atla; sistem agent part3_rca.analysis_branches verisinden deterministik doldurur. NEDEN 1 = ortak olay sorusu + A/B doğrudan neden cevabı; NEDEN 2–5 agent zinciri.
- Kök neden ve açıklama metinlerinde markdown kullanma: ** veya __ ile kalın vurgu yazma; düz Türkçe paragraf veya "1. Başlık:" gibi numaralı madde kullan.
- SADECE JSON döndür, başka hiçbir şey yazma
"""


# 
# DOCX YARDIMCI FONKSİYONLARI
# 

def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def _add_section_header(doc, number, title):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    _set_cell_bg(cell, COLOR["dark_blue"])
    _set_cell_margins(cell, 120, 120, 160, 160)
    p = cell.paragraphs[0]
    run = p.add_run(f"{number}. {title.upper()}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR["white"]
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def _add_subsection_header(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR["mid_blue"]


def _add_paragraph(doc, text, size=9, color=None, bold=False, italic=False,
                   space_before=4, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color or COLOR["dark_grey"]


def _add_colored_box(doc, title, content, bg_color, title_color=None):
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Table Grid'
    tc = table.cell(0, 0)
    _set_cell_bg(tc, bg_color)
    _set_cell_margins(tc, 80, 80, 140, 140)
    p = tc.paragraphs[0]
    run = p.add_run(str(title))
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = title_color or COLOR["white"]
    cc = table.cell(1, 0)
    _set_cell_bg(cc, COLOR["light_grey"])
    _set_cell_margins(cc, 100, 100, 140, 140)
    p = cc.paragraphs[0]
    run = p.add_run(str(content))
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR["dark_grey"]
    doc.add_paragraph()


def _add_info_table(doc, data: dict, header_color=None):
    if not data:
        return
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Table Grid'
    hc = header_color or COLOR["light_blue"]
    for i, (key, val) in enumerate(data.items()):
        row = table.rows[i]
        lc = row.cells[0]
        _set_cell_bg(lc, hc)
        _set_cell_margins(lc)
        run = lc.paragraphs[0].add_run(str(key))
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR["dark_blue"]
        rc = row.cells[1]
        bg = COLOR["white"] if i % 2 == 0 else COLOR["light_grey"]
        _set_cell_bg(rc, bg)
        _set_cell_margins(rc)
        run = rc.paragraphs[0].add_run(str(val))
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR["dark_grey"]
    doc.add_paragraph()


def _add_bullet_list(doc, items: list, color=None):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(str(item))
        run.font.size = Pt(10)
        run.font.color.rgb = color or COLOR["dark_grey"]


def _add_page_break(doc):
    doc.add_page_break()


# 
# RAPOR BÖLÜM FONKSİYONLARI
# 

def _build_cover(doc, cover: dict, lang_code: str = "tr"):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(cover.get("title", _label(lang_code, "cover_title")))
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = COLOR["dark_blue"]
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(cover.get("subtitle", _label(lang_code, "cover_subtitle")))
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR["mid_blue"]
    run.italic = True
    doc.add_paragraph()
    # Gizlilik banner
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    _set_cell_bg(cell, COLOR["red"])
    _set_cell_margins(cell, 120, 120, 200, 200)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(cover.get("confidentiality", _label(lang_code, "cover_confidentiality")))
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR["white"]
    doc.add_paragraph()
    _add_info_table(doc, {
        _label(lang_code, "ref_no"): cover.get("ref_no", "N/A"),
        _label(lang_code, "date"): cover.get("date", "N/A"),
        _label(lang_code, "location"): cover.get("location", "N/A"),
        _label(lang_code, "incident_type"): cover.get("incident_type", "N/A"),
    }, COLOR["dark_blue"])
    _add_colored_box(
        doc,
        _label(lang_code, "incident_summary"),
        strip_hse_codes(str(cover.get("incident_summary_short", "") or "")),
        COLOR["dark_blue"],
    )
    _add_page_break(doc)


def _build_executive_summary(doc, es: dict, root_causes: list):
    _add_section_header(doc, "1", _L("section_executive_summary"))
    _add_subsection_header(doc, f"1.1 {_L('subsection_incident_summary')}")
    for field in ["what_happened", "immediate_response"]:
        if es.get(field):
            _add_paragraph(doc, strip_hse_codes(str(es[field])), space_after=6)
    doc.add_paragraph()
    _add_subsection_header(doc, f"1.2 {_L('subsection_key_findings')}")
    for finding in es.get("key_findings", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"  {strip_hse_codes(str(finding))}")
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR["dark_grey"]
    doc.add_paragraph()
    _add_subsection_header(doc, f"1.3 {_L('subsection_immediate_actions')}")
    actions = es.get("immediate_actions", [])
    if actions:
        table = doc.add_table(rows=len(actions) + 1, cols=3)
        table.style = 'Table Grid'
        for j, h in enumerate([_L("th_immediate_action"), _L("th_responsible"), _L("th_status")]):
            c = table.rows[0].cells[j]
            _set_cell_bg(c, COLOR["dark_blue"])
            _set_cell_margins(c)
            run = c.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR["white"]
        for i, act in enumerate(actions):
            row = table.rows[i + 1]
            vals = [act.get("action", ""), act.get("responsible", ""), act.get("status", "")]
            for j, val in enumerate(vals):
                c = row.cells[j]
                _set_cell_bg(c, COLOR["light_grey"] if i % 2 == 0 else COLOR["white"])
                _set_cell_margins(c)
                run = c.paragraphs[0].add_run(str(val))
                run.font.size = Pt(10)
    _add_page_break(doc)


def _build_incident_details(doc, details: dict):
    _add_section_header(doc, "2", _L("section_incident_details"))
    _add_subsection_header(doc, f"2.1 {_L('subsection_info_table')}")
    _add_info_table(doc, details.get("info_table", {}))
    _add_subsection_header(doc, f"2.2 {_L('subsection_event_details')}")
    _add_info_table(doc, details.get("event_table", {}))
    doc.add_paragraph()
    _add_subsection_header(doc, f"2.3 {_L('subsection_timeline')}")
    timeline = details.get("timeline", [])
    if timeline:
        table = doc.add_table(rows=len(timeline) + 1, cols=2)
        table.style = 'Table Grid'
        for j, h in enumerate([_L("th_time"), _L("th_event")]):
            c = table.rows[0].cells[j]
            _set_cell_bg(c, COLOR["mid_blue"])
            _set_cell_margins(c)
            run = c.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR["white"]
        for i, step in enumerate(timeline):
            row = table.rows[i + 1]
            bg = COLOR["light_blue"] if i % 2 == 0 else COLOR["white"]
            tc = row.cells[0]
            _set_cell_bg(tc, bg)
            _set_cell_margins(tc)
            run = tc.paragraphs[0].add_run(step.get("time", ""))
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR["dark_blue"]
            ec = row.cells[1]
            _set_cell_bg(ec, bg)
            _set_cell_margins(ec)
            run = ec.paragraphs[0].add_run(step.get("event", ""))
            run.font.size = Pt(10)
    doc.add_paragraph()
    sev = details.get("severity", {})
    if sev:
        _add_subsection_header(doc, f"2.4 {_L('subsection_severity')}")
        _add_info_table(doc, {
            _L("severity_actual"): sev.get("actual_harm", ""),
            _L("severity_potential"): sev.get("potential_harm", ""),
            _L("severity_investigation"): sev.get("investigation_level", ""),
        })
    _add_page_break(doc)


def _build_analysis_method(doc, method: dict):
    method = _normalize_analysis_method(method)
    _add_section_header(doc, "3", _L("section_analysis_method"))
    _add_subsection_header(doc, f"3.1 {_L('subsection_five_why')}")
    _add_paragraph(doc, method.get("five_why_explanation", ""), space_after=8)
    _add_subsection_header(doc, f"3.2 {_L('subsection_factor_categories')}")
    codes = method.get("code_system", [])
    if codes:
        table = doc.add_table(rows=len(codes) + 1, cols=2)
        table.style = 'Table Grid'
        for j, h in enumerate([_L("th_category"), _L("th_description")]):
            c = table.rows[0].cells[j]
            _set_cell_bg(c, COLOR["dark_blue"])
            _set_cell_margins(c)
            run = c.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR["white"]
        for i, code in enumerate(codes):
            row = table.rows[i + 1]
            bg = COLOR["light_grey"] if i % 2 == 0 else COLOR["white"]
            vals = [code.get("category",""), code.get("description","")]
            for j, val in enumerate(vals):
                c = row.cells[j]
                _set_cell_bg(c, bg)
                _set_cell_margins(c)
                run = c.paragraphs[0].add_run(str(val))
                run.bold = (j == 0)
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR["dark_blue"] if j == 0 else COLOR["dark_grey"]
    doc.add_paragraph()
    _add_subsection_header(doc, f"3.3 {_L('subsection_analysis_team')}")
    members = method.get("team_members", [])
    if members:
        table = doc.add_table(rows=len(members) + 1, cols=3)
        table.style = 'Table Grid'
        for j, h in enumerate([_L("th_name"), _L("th_role"), _L("th_date")]):
            c = table.rows[0].cells[j]
            _set_cell_bg(c, COLOR["mid_blue"])
            _set_cell_margins(c)
            run = c.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR["white"]
        for i, m in enumerate(members):
            row = table.rows[i + 1]
            bg = COLOR["light_grey"] if i % 2 == 0 else COLOR["white"]
            for j, val in enumerate([m.get("name",""), m.get("role",""), m.get("date","")]):
                c = row.cells[j]
                _set_cell_bg(c, bg)
                _set_cell_margins(c)
                run = c.paragraphs[0].add_run(str(val))
                run.font.size = Pt(10)
    _add_page_break(doc)


def _build_branches(doc, branches: list):
    branch_colors = [COLOR["red"], COLOR["orange"], COLOR["green"], COLOR["mid_blue"]]
    for branch in branches:
        bn = branch.get("branch_number", 1)
        color = branch_colors[(bn - 1) % len(branch_colors)]
        # Dal başlığı
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        _set_cell_bg(cell, color)
        _set_cell_margins(cell, 120, 120, 160, 160)
        p = cell.paragraphs[0]
        run = p.add_run(branch.get("branch_title", f"{_L('branch_critical_factor')} {bn}"))
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR["white"]
        doc.add_paragraph()
        _add_subsection_header(doc, f"{3+bn}.1 {_L('subsection_why_table')}")
        why_chain = branch.get("why_chain", [])
        if why_chain:
            table = doc.add_table(rows=len(why_chain) + 1, cols=2)
            table.style = 'Table Grid'
            for j, h in enumerate([_L("why_number_col"), _L("why_qa_col")]):
                c = table.rows[0].cells[j]
                _set_cell_bg(c, COLOR["dark_blue"])
                _set_cell_margins(c)
                run = c.paragraphs[0].add_run(h)
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR["white"]
            for i, why in enumerate(why_chain):
                row = table.rows[i + 1]
                bg = COLOR["light_grey"] if i % 2 == 0 else COLOR["white"]
                q = strip_hse_codes(str(why.get("question", "") or why.get("question_tr", "") or ""))
                a = strip_hse_codes(str(why.get("answer", "") or why.get("answer_tr", "") or ""))
                qa = f"{_L('why_prefix')}: {q}\n{_L('answer_prefix')}: {a}"
                wn = why.get('number') or why.get('level') or (i + 1)
                vals = [f"{_L('why_prefix')} {wn}", qa]
                for j, val in enumerate(vals):
                    c = row.cells[j]
                    _set_cell_bg(c, bg)
                    _set_cell_margins(c)
                    run = c.paragraphs[0].add_run(str(val))
                    run.font.size = Pt(9)
                    run.bold = (j == 0)
        doc.add_paragraph()
        _add_subsection_header(doc, f"{3+bn}.2 {_L('subsection_root_cause')}")
        rc_label = strip_root_cause_label_prefix(
            str(branch.get("root_cause_title", "") or ""),
            branch_number=bn,
        )
        rc_title = f"{_L('root_cause_prefix')} {bn}: {rc_label}"
        rc_content = strip_hse_codes(str(branch.get('root_cause_detail','') or ''))
        _add_colored_box(doc, rc_title, rc_content, color)
        org_factors = branch.get("organizational_factors", [])
        if org_factors:
            _add_subsection_header(doc, f"{3+bn}.3 {_L('subsection_org_factors')}")
            _add_bullet_list(doc, [strip_hse_codes(str(x)) for x in org_factors])
        _add_page_break(doc)


def _build_meta_root_cause(doc, meta: dict):
    """Meta kök neden bölümü oluşturur (tüm dalların ortak paydası)"""
    if not meta or not meta.get("exists"):
        return
    
    _add_section_header(doc, "5", _L("section_meta_root_cause"))
    _add_subsection_header(doc, f"5.1 {_L('subsection_systemic_weakness')}")
    
    # Ana meta kök neden kutusu
    meta_title = strip_hse_codes(str(meta.get('title', 'Meta Kök Neden') or ''))
    meta_desc = strip_hse_codes(str(meta.get('description', '') or ''))
    _add_colored_box(doc, meta_title, meta_desc, COLOR["red"], COLOR["white"])
    
    # Sentezlenen kodlar
    synthesized = meta.get('synthesized_from', [])
    if synthesized:
        _add_subsection_header(doc, "5.2 Sentezlenen Kök Nedenler")
        synth_text = (
            f"Bu meta kök neden, analizdeki {len(synthesized)} dalın ortak paydasından türetilmiş "
            "üst düzey bir sistemik zayıflığı ifade eder."
        )
        _add_paragraph(doc, synth_text, space_after=8)
    
    # Sistemik zayıflık
    if meta.get('systemic_weakness'):
        _add_subsection_header(doc, "5.3 Sistemik Zayıflık")
        _add_paragraph(doc, strip_hse_codes(str(meta.get('systemic_weakness', '') or '')), space_after=8)
    
    # Stratejik sonuçlar
    implications = meta.get('strategic_implications', [])
    if implications:
        _add_subsection_header(doc, "5.4 Stratejik Sonuçlar ve Etkiler")
        _add_bullet_list(doc, implications, COLOR["red"])
    
    doc.add_paragraph()
    _add_page_break(doc)


def _build_contributing_factors(doc, factors: list):
    _add_section_header(doc, "6", _L("section_contributing_short"))
    doc.add_paragraph()
    priority_colors = {"Yüksek": COLOR["red"], "Orta": COLOR["orange"], "Düşük": COLOR["green"],
                       "High": COLOR["red"], "Medium": COLOR["orange"], "Low": COLOR["green"]}
    if factors:
        table = doc.add_table(rows=len(factors) + 1, cols=3)
        table.style = 'Table Grid'
        for j, h in enumerate([_L("th_factor_type"), _L("th_description"), _L("th_impact_level")]):
            c = table.rows[0].cells[j]
            _set_cell_bg(c, COLOR["dark_blue"])
            _set_cell_margins(c)
            run = c.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR["white"]
        for i, f in enumerate(factors):
            row = table.rows[i + 1]
            bg = COLOR["light_grey"] if i % 2 == 0 else COLOR["white"]
            impact = f.get("impact_level","Orta")
            for j, val in enumerate([f.get("factor_type",""), f.get("description",""), impact]):
                c = row.cells[j]
                if j == 2:
                    _set_cell_bg(c, priority_colors.get(impact, COLOR["light_grey"]))
                    run = c.paragraphs[0].add_run(str(val))
                    run.font.color.rgb = COLOR["white"]
                else:
                    _set_cell_bg(c, bg)
                    run = c.paragraphs[0].add_run(str(val))
                    run.font.color.rgb = COLOR["dark_grey"]
                _set_cell_margins(c)
                run.bold = (j == 0)
                run.font.size = Pt(10)
    _add_page_break(doc)


def _build_corrective_actions(doc, actions: list):
    _add_section_header(doc, "7", _L("section_corrective_actions"))
    doc.add_paragraph()
    priority_colors = {
        "ACİL": COLOR["red"], "YÜKSEK": COLOR["orange"],
        "ORTA": COLOR["green"], "DÜŞÜK": COLOR["mid_blue"],
        "URGENT": COLOR["red"], "HIGH": COLOR["orange"],
        "MEDIUM": COLOR["green"], "LOW": COLOR["mid_blue"],
    }
    if actions:
        table = doc.add_table(rows=len(actions) + 1, cols=6)
        table.style = 'Table Grid'
        for j, h in enumerate([_L("th_no"), _L("th_activity"), _L("th_priority"), _L("th_responsible"), _L("th_duration"), _L("th_kpi")]):
            c = table.rows[0].cells[j]
            _set_cell_bg(c, COLOR["dark_blue"])
            _set_cell_margins(c)
            run = c.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR["white"]
        for i, act in enumerate(actions):
            row = table.rows[i + 1]
            bg = COLOR["light_grey"] if i % 2 == 0 else COLOR["white"]
            priority = act.get("priority","ORTA")
            vals = [str(act.get("no",i+1)), act.get("action",""), priority,
                    act.get("responsible",""), act.get("deadline",""), act.get("kpi","")]
            for j, val in enumerate(vals):
                c = row.cells[j]
                if j == 2:
                    _set_cell_bg(c, priority_colors.get(priority, COLOR["light_grey"]))
                    run = c.paragraphs[0].add_run(str(val))
                    run.font.color.rgb = COLOR["white"]
                    run.bold = True
                else:
                    _set_cell_bg(c, bg)
                    run = c.paragraphs[0].add_run(str(val))
                    run.font.color.rgb = COLOR["dark_grey"]
                _set_cell_margins(c)
                run.font.size = Pt(9)
    _add_page_break(doc)


def _build_lessons_learned(doc, lessons: dict):
    _add_section_header(doc, "9", _L("section_lessons_learned"))
    doc.add_paragraph()
    sections = [
        (_L("lesson_what_to_do"), lessons.get("what_to_do", []), COLOR["green"]),
        (_L("lesson_long_term"), lessons.get("long_term", []), COLOR["mid_blue"]),
        (_L("lesson_communication"), lessons.get("communication", []), COLOR["orange"]),
        (_L("lesson_training"), lessons.get("training", []), COLOR["red"]),
    ]
    for title, items, color in sections:
        if items:
            _add_colored_box(doc, title, "\n".join(f"- {item}" for item in items), color)
    _add_page_break(doc)


def _build_conclusion(doc, conclusion: dict):
    _add_section_header(doc, "10", _L("section_conclusion"))
    _add_subsection_header(doc, f"10.1 {_L('subsection_general_assessment')}")
    _add_paragraph(doc, conclusion.get("overall_assessment",""), space_after=8)
    _add_subsection_header(doc, f"10.2 {_L('subsection_short_term')}")
    _add_bullet_list(doc, conclusion.get("short_term_measures",[]))
    doc.add_paragraph()
    _add_subsection_header(doc, f"10.3 {_L('subsection_long_term')}")
    _add_bullet_list(doc, conclusion.get("long_term_improvements",[]))
    doc.add_paragraph()
    _add_subsection_header(doc, f"10.4 {_L('subsection_comparison')}")
    comparison = conclusion.get("comparison_table", [])
    if comparison:
        table = doc.add_table(rows=len(comparison) + 1, cols=3)
        table.style = 'Table Grid'
        for j, h in enumerate([_L("th_criterion"), _L("th_current"), _L("th_target")]):
            c = table.rows[0].cells[j]
            _set_cell_bg(c, COLOR["dark_blue"])
            _set_cell_margins(c)
            run = c.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR["white"]
        for i, row_data in enumerate(comparison):
            row = table.rows[i + 1]
            vals = [row_data.get("criterion",""), row_data.get("current",""), row_data.get("target","")]
            for j, val in enumerate(vals):
                c = row.cells[j]
                if j == 2:
                    _set_cell_bg(c, COLOR["green"])
                    run = c.paragraphs[0].add_run(str(val))
                    run.font.color.rgb = COLOR["white"]
                elif j == 1:
                    _set_cell_bg(c, COLOR["light_grey"])
                    run = c.paragraphs[0].add_run(str(val))
                    run.font.color.rgb = COLOR["red"]
                else:
                    _set_cell_bg(c, COLOR["light_blue"])
                    run = c.paragraphs[0].add_run(str(val))
                    run.font.color.rgb = COLOR["dark_blue"]
                _set_cell_margins(c)
                run.bold = (j == 0)
                run.font.size = Pt(10)
    _add_page_break(doc)


def _build_signature_page(doc):
    _add_section_header(doc, "11", _L("section_signatures"))
    doc.add_paragraph()
    roles = [
        (_L("sig_role_prepared"), "HSE Uzmanı", "HSE Kök Neden Analisti"),
        (_L("sig_role_reviewed"), "HSE Yöneticisi", "HSE Departman Yöneticisi"),
        (_L("sig_role_approved"), "Tesis Müdürü", "Genel Operasyon Müdürü"),
    ]
    table = doc.add_table(rows=len(roles) + 1, cols=4)
    table.style = 'Table Grid'
    for j, h in enumerate([_L("th_role"), _L("th_name"), _L("th_title"), _L("th_signature")]):
        c = table.rows[0].cells[j]
        _set_cell_bg(c, COLOR["dark_blue"])
        _set_cell_margins(c)
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR["white"]
    for i, (role, name, title) in enumerate(roles):
        row = table.rows[i + 1]
        bg = COLOR["light_grey"] if i % 2 == 0 else COLOR["white"]
        for j, val in enumerate([role, name, title, "___________________\n_____ / _____ / _____"]):
            c = row.cells[j]
            _set_cell_bg(c, bg)
            _set_cell_margins(c, 160, 160, 120, 120)
            run = c.paragraphs[0].add_run(str(val))
            run.bold = (j == 0)
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR["dark_blue"] if j == 0 else COLOR["dark_grey"]
    
    # Disclaimer
    doc.add_paragraph()
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run("---")
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR["dark_grey"]
    
    doc.add_paragraph()
    disclaimer_text = doc.add_paragraph()
    disclaimer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer_text.add_run(_L("ai_disclaimer"))
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR["dark_grey"]
    run.italic = True


# 
# DİL TESPİTİ
# 

def detect_language(text: str) -> dict:
    """
    Ham olay metninden dili tespit eder.
    Returns dict with: code, name, rtl, html_lang, font_hint
    """
    if not text:
        return {"code": "tr", "name": "Turkish", "rtl": False, "html_lang": "tr", "font_hint": ""}

    # Unicode blok örnekleme (ilk 500 karakter yeterli)
    sample = text[:500]

    arabic_chars  = sum(1 for c in sample if "\u0600" <= c <= "\u06FF")
    hebrew_chars  = sum(1 for c in sample if "\u0590" <= c <= "\u05FF")
    cyrillic_chars = sum(1 for c in sample if "\u0400" <= c <= "\u04FF")
    cjk_chars     = sum(1 for c in sample if "\u4E00" <= c <= "\u9FFF")
    latin_chars   = sum(1 for c in sample if c.isalpha() and ord(c) < 0x250)

    total = max(len([c for c in sample if c.isalpha()]), 1)

    # Arapça / Farsça / Urduca
    if arabic_chars / total > 0.15:
        return {"code": "ar", "name": "Arabic", "rtl": True,
                "html_lang": "ar", "font_hint": "Arial Unicode MS"}

    # İbranice
    if hebrew_chars / total > 0.15:
        return {"code": "he", "name": "Hebrew", "rtl": True,
                "html_lang": "he", "font_hint": "Arial Unicode MS"}

    # Kiril (Rusça/Bulgarca)
    if cyrillic_chars / total > 0.15:
        return {"code": "ru", "name": "Russian", "rtl": False,
                "html_lang": "ru", "font_hint": ""}

    # Çince/Japonca/Korece
    if cjk_chars / total > 0.10:
        return {"code": "zh", "name": "Chinese", "rtl": False,
                "html_lang": "zh", "font_hint": ""}

    # Latin alfabesi — İngilizce vs Türkçe vs İspanyolca vs Fransızca
    if latin_chars / total > 0.50:
        lower = sample.lower()
        # Türkçe karakterler
        turkish_chars = sum(1 for c in lower if c in "şğüıöçŞĞÜİÖÇ")
        # İspanyolca karakterler
        spanish_chars = sum(1 for c in lower if c in "áéíóúüñ¿¡")
        # Fransızca karakterler (è, ê, î, ô, û, œ, æ, ç — İspanyolca ile çakışmayan)
        french_chars = sum(1 for c in lower if c in "èêîôûœæàù")
        # Portekizce
        portuguese_chars = sum(1 for c in lower if c in "ãõàâêîôûçé")

        # Yaygın İngilizce kelimeler
        en_words = {"the", "and", "was", "not", "with", "from", "work", "site",
                    "worker", "incident", "fall", "height", "safety", "injury"}
        tr_words = {"bir", "ve", "bu", "ile", "için", "çalışan", "kaza", "olay",
                    "iş", "işçi", "güvenlik"}
        # İspanyolcaya özgü — Fransızca'da olmayan kelimeler
        es_words = {"el", "los", "una", "del", "que", "trabajador", "accidente",
                    "seguridad", "caída", "andamio", "herido", "altura"}
        # Fransızcaya özgü — İspanyolca'da olmayan kelimeler
        fr_words = {"les", "des", "dans", "sur", "avec", "un", "est",
                    "travailleur", "accident", "sécurité", "tombé", "chute",
                    "ouvrier", "bâtiment", "grièvement", "portait", "harnais"}

        words = set(lower.split())
        en_score = len(words & en_words)
        tr_score = len(words & tr_words) + turkish_chars * 3
        es_score = len(words & es_words) + spanish_chars * 3
        fr_score = len(words & fr_words) + french_chars * 3

        scores = {"en": en_score, "tr": tr_score, "es": es_score, "fr": fr_score}
        best = max(scores, key=lambda k: scores[k])

        lang_map = {
            "en": {"code": "en", "name": "English",  "rtl": False, "html_lang": "en", "font_hint": ""},
            "tr": {"code": "tr", "name": "Turkish",  "rtl": False, "html_lang": "tr", "font_hint": ""},
            "es": {"code": "es", "name": "Spanish",  "rtl": False, "html_lang": "es", "font_hint": ""},
            "fr": {"code": "fr", "name": "French",   "rtl": False, "html_lang": "fr", "font_hint": ""},
        }
        return lang_map.get(best, lang_map["en"])

    # Fallback
    return {"code": "tr", "name": "Turkish", "rtl": False, "html_lang": "tr", "font_hint": ""}


# 
# ANA AGENT SINIFI
# 

class SkillBasedDocxAgent:
    """
    V2: OpenRouter uzerinden LLM icerigi uretir → python-docx DOCX olusturur.

    Kullanım:
        agent = SkillBasedDocxAgent()
        path = agent.generate_report(
            investigation_data=data,
            output_path="outputs/rapor.docx"
        )
    """

    def __init__(self, api_key: Optional[str] = None):
        load_dotenv()
        key = api_key or os.getenv("OPENROUTER_API_KEY")
        if not key:
            raise ValueError("OPENROUTER_API_KEY bulunamadı! .env dosyasına ekleyin.")
        self.api_key = key
        self.model = resolve_openrouter_docx_model().strip()
        self.api_url = _resolve_openrouter_chat_completions_url()
        print(f" SkillBasedDocxAgent V2 hazır (OpenRouter {self.model})")
        print(f"   API: {self.api_url}")

    def generate_report(
        self,
        investigation_data: Dict,
        output_path: str = "outputs/hse_report.docx",
        timeout_seconds: int = 600,
        preferred_language: str = "",
    ) -> str:
        """
        Investigation data'dan kapsamlı DOCX rapor üretir.

        Args:
            investigation_data: part1, part2, part3_rca içeren tam pipeline verisi
            output_path: Çıktı dosyası yolu
            timeout_seconds: API timeout (saniye)

        Returns:
            Oluşturulan DOCX dosyasının tam yolu
        """
        print("\n" + "=" * 70)
        print(f" DOCX RAPOR URETME V2 (OpenRouter model: {self.model})")
        print("=" * 70)

        layout = resolve_report_layout(investigation_data)
        investigation_data = {**investigation_data, "report_layout": layout}
        set_report_text_policy(show_technical_codes=bool(layout.get("show_technical_codes")))
        try:
            return self._generate_report_impl(
                investigation_data,
                output_path,
                timeout_seconds,
                preferred_language,
            )
        finally:
            set_report_text_policy(show_technical_codes=False)

    def _generate_report_impl(
        self,
        investigation_data: Dict,
        output_path: str,
        timeout_seconds: int,
        preferred_language: str,
    ) -> str:
        raw_data = self._build_raw_payload(investigation_data)

        #  Dil tespiti 
        source_text = (
            raw_data.get("part3_rca", {}).get("incident_summary", "")
            or raw_data.get("part1", {}).get("description", "")
            or json.dumps(raw_data, ensure_ascii=False)[:800]
        )
        preferred = (
            (preferred_language or "").strip().lower()
            or str(investigation_data.get("output_language") or "").strip().lower()
            or str(investigation_data.get("preferred_language") or "").strip().lower()
        )
        forced_lang_map = {
            "tr": {"code": "tr", "name": "Turkish", "rtl": False, "html_lang": "tr", "font_hint": ""},
            "en": {"code": "en", "name": "English", "rtl": False, "html_lang": "en", "font_hint": ""},
            "de": {"code": "de", "name": "German", "rtl": False, "html_lang": "de", "font_hint": ""},
            "fr": {"code": "fr", "name": "French", "rtl": False, "html_lang": "fr", "font_hint": ""},
            "es": {"code": "es", "name": "Spanish", "rtl": False, "html_lang": "es", "font_hint": ""},
            "ar": {"code": "ar", "name": "Arabic", "rtl": True, "html_lang": "ar", "font_hint": "Arial Unicode MS"},
        }
        lang = forced_lang_map.get(preferred) or detect_language(source_text)
        set_report_lang(lang["code"])
        print(f" Tespit edilen dil: {lang['name']} ({lang['code']}) | RTL: {lang['rtl']}")
        # 

        char_count = len(json.dumps(raw_data, ensure_ascii=False))
        print(f" Ham veri hazır ({char_count} karakter)")

        print(f"\n OpenRouter modeline icerik istegi gonderiliyor... ({self.model})")
        start = time.time()
        content = self._generate_content_with_claude(raw_data, lang)
        elapsed = time.time() - start
        out_chars = len(json.dumps(content, ensure_ascii=False))
        print(f" İçerik alındı ({elapsed:.1f}s, {out_chars} karakter)")

        print("\n DOCX oluşturuluyor (python-docx)...")
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        self._build_docx(content, str(output_file.resolve()), lang)

        if not output_file.exists():
            raise RuntimeError(f"DOCX oluşturulamadı: {output_file}")

        size_kb = output_file.stat().st_size / 1024
        print(f"\n DOCX başarıyla oluşturuldu!")
        print(f" Dosya : {output_file.resolve()}")
        print(f" Boyut : {size_kb:.1f} KB")
        
        # HTML rapor da üret
        html_path = str(output_file).replace('.docx', '.html')
        print(f"\n HTML raporu oluşturuluyor...")
        self._build_html(content, html_path, lang, investigation_data)
        html_size_kb = Path(html_path).stat().st_size / 1024
        print(f" HTML başarıyla oluşturuldu!")
        print(f" Dosya : {html_path}")
        print(f" Boyut : {html_size_kb:.1f} KB")
        
        # 5-Why Decision Tree HTML'ini oluştur
        decision_tree_path = str(output_file).replace('.docx', '_decision_tree.html')
        print(f"\n 5-Why Decision Tree oluşturuluyor...")
        self._build_decision_tree(investigation_data, decision_tree_path)
        if Path(decision_tree_path).exists():
            dt_size_kb = Path(decision_tree_path).stat().st_size / 1024
            print(f" Decision Tree başarıyla oluşturuldu!")
            print(f" Dosya : {decision_tree_path}")
            print(f" Boyut : {dt_size_kb:.1f} KB")
        
        print("=" * 70)
        return str(output_file.resolve())

    def _build_raw_payload(self, data: Dict) -> Dict:
        if "part3_rca" in data:
            return {
                "part1": data.get("part1", {}),
                "part2": data.get("part2", {}),
                "part3_rca": data["part3_rca"],
            }
        if "analysis_branches" in data:
            return {"part1": {}, "part2": {}, "part3_rca": data}
        return data

    @staticmethod
    def _is_empty_value(v: Any) -> bool:
        if v is None:
            return True
        if isinstance(v, str):
            return not v.strip()
        if isinstance(v, (list, tuple, set, dict)):
            return len(v) == 0
        return False

    def _merge_with_fallback_content(self, parsed: Dict, fallback: Dict) -> Dict:
        """LLM çıktısı kısmi/boşsa fallback'ten doldur; dolu alanları koru."""
        if not isinstance(parsed, dict):
            return fallback
        merged = copy.deepcopy(parsed)
        for key, fb_val in fallback.items():
            cur_val = merged.get(key)
            if key not in merged or self._is_empty_value(cur_val):
                merged[key] = copy.deepcopy(fb_val)
                continue
            if isinstance(cur_val, dict) and isinstance(fb_val, dict):
                merged[key] = self._merge_with_fallback_content(cur_val, fb_val)
        return merged

    def _enforce_taxonomy_titles(self, content: Dict, raw_data: Dict) -> Dict:
        """LLM kısaltmalarını ezip BARSEL tablo başlıklarını (kodsuz) rapora yazar."""
        if not isinstance(content, dict):
            return content
        try:
            from agents.barsel_taxonomy import (
                apply_official_taxonomy_titles_to_report_branches,
                apply_official_taxonomy_titles_to_root_causes,
            )
        except ImportError:
            from .barsel_taxonomy import (
                apply_official_taxonomy_titles_to_report_branches,
                apply_official_taxonomy_titles_to_root_causes,
            )
        part3 = raw_data.get("part3_rca") or {}
        raw_branches = part3.get("analysis_branches") or part3.get("branches") or []
        out = copy.deepcopy(content)
        if out.get("branches"):
            out["branches"] = apply_official_taxonomy_titles_to_report_branches(
                out["branches"],
                raw_branches,
            )
        if out.get("root_causes"):
            out["root_causes"] = apply_official_taxonomy_titles_to_root_causes(
                out["root_causes"],
                raw_branches,
            )
        return out

    def _pin_agent_why_chains(self, content: Dict, raw_data: Dict) -> Dict:
        """P1.26 R1: Merge sonrası why_chain'i agent part3_rca verisine sabitle."""
        try:
            from agents.report_why_chain import pin_agent_why_chains_to_report
        except ImportError:
            from .report_why_chain import pin_agent_why_chains_to_report
        return pin_agent_why_chains_to_report(content, raw_data)

    def _build_deterministic_fallback_content(self, raw_data: Dict, lang: Optional[Dict] = None) -> Dict:
        """LLM başarısız olsa bile boş olmayan rapor iskeleti üret."""
        lang = lang or {"code": "tr"}
        lang_code = (lang.get("code") or "tr").lower()
        part1 = raw_data.get("part1", {}) or {}
        part2 = raw_data.get("part2", {}) or {}
        part3 = raw_data.get("part3_rca", {}) or {}
        overview = part1.get("overview", {}) if isinstance(part1.get("overview", {}), dict) else {}

        incident_summary = str(
            part3.get("incident_summary")
            or overview.get("what_happened")
            or part1.get("description")
            or "Olay özeti mevcut değil."
        )
        incident_id = str(
            part1.get("incident_id")
            or part1.get("reference_no")
            or part3.get("incident_id")
            or "N/A"
        )
        location = str(
            part1.get("location")
            or overview.get("where_happened")
            or "N/A"
        )
        incident_type = str(
            part1.get("incident_type")
            or part2.get("incident_type")
            or "N/A"
        )

        branches_raw = part3.get("analysis_branches") or part3.get("branches") or []
        branches: List[Dict[str, Any]] = []
        root_causes: List[Dict[str, Any]] = []
        shared_event_q = ""
        try:
            from agents.report_why_chain import build_shared_event_question

            shared_event_q = build_shared_event_question(incident_summary)
        except ImportError:
            from .report_why_chain import build_shared_event_question

            shared_event_q = build_shared_event_question(incident_summary)

        for idx, br in enumerate(branches_raw[:8], start=1):
            immediate = br.get("immediate_cause", {}) if isinstance(br.get("immediate_cause", {}), dict) else {}
            root = br.get("root_cause", {}) if isinstance(br.get("root_cause", {}), dict) else {}
            try:
                from agents.report_why_chain import build_pinned_why_chain
            except ImportError:
                from .report_why_chain import build_pinned_why_chain
            why_chain, _ = build_pinned_why_chain(br, shared_event_q)

            root_code = str(root.get("code") or "").strip().upper()
            try:
                from agents.barsel_taxonomy import (
                    critical_factor_title_for_code,
                    enrich_root_cause_from_taxonomy,
                    extract_taxonomy_code,
                    resolve_root_cause_code_from_branch,
                    root_cause_leaf_title_for_code,
                )
            except ImportError:
                from .barsel_taxonomy import (
                    critical_factor_title_for_code,
                    enrich_root_cause_from_taxonomy,
                    extract_taxonomy_code,
                    resolve_root_cause_code_from_branch,
                    root_cause_leaf_title_for_code,
                )
            root_enriched = enrich_root_cause_from_taxonomy(
                root,
                incident_hint=str(immediate.get("evidence_tr") or incident_summary[:400]),
            )
            root_code = (
                resolve_root_cause_code_from_branch(br)
                or extract_taxonomy_code(str(root_enriched.get("code") or root_code or ""))
            )
            cf_title = critical_factor_title_for_code(root_code)
            root_title = strip_root_cause_label_prefix(
                root_cause_leaf_title_for_code(root_code)
                or str(root_enriched.get("standard_title_tr") or "").strip(),
                branch_number=idx,
            )
            root_section = ""
            root_detail = sanitize_report_text(
                str(root_enriched.get("explanation_tr") or root_enriched.get("explanation") or root_title)
            )
            branch_title = f"KRİTİK FAKTÖR {idx}"
            if cf_title:
                branch_title = f"KRİTİK FAKTÖR {idx} - {strip_hse_codes(cf_title)}"
            branches.append(
                {
                    "branch_number": br.get("branch_number", idx),
                    "branch_title": branch_title,
                    "initial_condition": strip_hse_codes(str(immediate.get("evidence_tr") or incident_summary[:500])),
                    "direct_cause": _direct_cause_sentence(
                        strip_hse_codes(str(immediate.get("cause_tr") or immediate.get("cause") or "")),
                        lang_code,
                    ),
                    "why_chain": why_chain,
                    "root_cause_title": strip_hse_codes(root_title),
                    "root_cause_section": strip_hse_codes(root_section),
                    "root_cause_detail": root_detail,
                    "organizational_factors": [],
                }
            )
            root_causes.append(
                {
                    "title": root_title,
                    "section": root_section,
                    "category": str(root.get("category_type") or root.get("category") or ""),
                    "contributing_organizations": "",
                    "detailed_description": root_detail,
                    "impacts": [],
                }
            )

        # analysis_branches yoksa part3.root_causes listesinden üret
        if not root_causes:
            for idx, rc in enumerate(part3.get("root_causes", [])[:8], start=1):
                rc_code = str(rc.get("code") or "").strip().upper()
                try:
                    from agents.barsel_taxonomy import (
                        critical_factor_title_for_code,
                        enrich_root_cause_from_taxonomy,
                        extract_taxonomy_code,
                        root_cause_leaf_title_for_code,
                    )
                except ImportError:
                    from .barsel_taxonomy import (
                        critical_factor_title_for_code,
                        enrich_root_cause_from_taxonomy,
                        extract_taxonomy_code,
                        root_cause_leaf_title_for_code,
                    )
                rc_enriched = enrich_root_cause_from_taxonomy(rc)
                rc_code = extract_taxonomy_code(
                    str(rc_enriched.get("code") or rc_code or "")
                )
                cf_title = critical_factor_title_for_code(rc_code)
                title = strip_root_cause_label_prefix(
                    root_cause_leaf_title_for_code(rc_code)
                    or str(rc_enriched.get("standard_title_tr") or "").strip(),
                    branch_number=idx,
                )
                detail = strip_hse_codes(
                    str(rc_enriched.get("explanation_tr") or rc_enriched.get("explanation") or title)
                )
                root_section = ""
                branch_title = f"KRİTİK FAKTÖR {idx}"
                if cf_title:
                    branch_title = f"KRİTİK FAKTÖR {idx} - {strip_hse_codes(cf_title)}"
                root_causes.append(
                    {
                        "title": strip_hse_codes(title),
                        "section": strip_hse_codes(root_section),
                        "category": str(rc.get("category") or rc.get("category_type") or ""),
                        "contributing_organizations": "",
                        "detailed_description": detail,
                        "impacts": [],
                    }
                )
                branches.append(
                    {
                        "branch_number": idx,
                        "branch_title": branch_title,
                        "initial_condition": strip_hse_codes(incident_summary[:500]),
                        "direct_cause": _direct_cause_sentence(title, lang_code),
                        "why_chain": [],
                        "root_cause_title": title,
                        "root_cause_section": root_section,
                        "root_cause_detail": detail,
                        "organizational_factors": [],
                    }
                )

        immediate_actions = [
            {"action": "Tehlikeli iş adımı durduruldu ve alan emniyete alındı", "responsible": "Saha Sorumlusu", "status": "Tamamlandı"},
            {"action": "İlgili ekip için hızlı güvenlik bilgilendirmesi yapıldı", "responsible": "HSE", "status": "Devam ediyor"},
        ]

        return {
            "cover": {
                "title": _label(lang_code, "cover_title"),
                "subtitle": _label(lang_code, "cover_subtitle"),
                "confidentiality": _label(lang_code, "cover_confidentiality"),
                "ref_no": incident_id,
                "date": datetime.now().strftime("%Y-%m-%d"),
                "location": location,
                "incident_type": incident_type,
                "incident_summary_short": short_incident_summary(
                    strip_hse_codes(incident_summary), 360
                ),
            },
            "executive_summary": {
                "what_happened": strip_hse_codes(incident_summary),
                "immediate_response": strip_hse_codes(str(part1.get("immediate_response") or part2.get("immediate_response") or "")),
                "key_findings": [rc.get("title", "") for rc in root_causes[:5] if rc.get("title")],
                "immediate_actions": immediate_actions,
            },
            "incident_details": {
                "info_table": {
                    "Referans No": incident_id,
                    "Tarih": datetime.now().strftime("%Y-%m-%d"),
                    "Lokasyon": location,
                    "Olay Tipi": incident_type,
                },
                "event_table": {
                    "Analiz Seviyesi": str(part2.get("investigation_level") or ""),
                },
                "event_summary_only": short_incident_summary(
                    strip_hse_codes(incident_summary), 360
                ),
                "timeline": [],
            },
            "analysis_method": {
                "five_why_explanation": "HSG245 tabanlı 5-Why yaklaşımı ile doğrudan ve kök nedenler analiz edilmiştir.",
                "code_system": _default_barsel_code_system(get_report_lang()),
                "team_members": [],
            },
            "branches": branches,
            "root_causes": root_causes,
            "contributing_factors": [],
            "corrective_actions": [],
            "lessons_learned": {
                "what_to_do": [],
                "long_term": [],
                "communication": [],
                "training": [],
            },
            "conclusion": {
                "overall_assessment": "Rapor, mevcut vaka verileri üzerinden otomatik olarak derlenmiştir.",
                "short_term_measures": [],
                "long_term_improvements": [],
                "comparison_table": [],
            },
        }

    def _generate_content_with_claude(self, raw_data: Dict, lang: Optional[Dict] = None) -> Dict:
        lang = lang or {"code": "tr", "name": "Turkish", "rtl": False}
        lang_name = lang["name"]
        lang_instruction = (
            f"\n\nCRITICAL LANGUAGE RULE: The incident data is written in {lang_name}. "
            f"You MUST write ALL report content ENTIRELY in {lang_name}. "
            f"Do NOT mix languages. JSON keys stay in English but every VALUE must be in {lang_name}.\n"
        )
        user_msg = (
            lang_instruction
            + "Aşağıdaki kök neden analizi ham verisini kullanarak "
            "profesyonel HSE raporu içeriğini üret.\n\n"
            "Ham Veri:\n```json\n"
            + json.dumps(raw_data, ensure_ascii=False, indent=2)
            + "\n```\n\n"
            "SADECE JSON döndür. Başka hiçbir şey yazma."
        )

        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://github.com/hse-rca-system",
            "X-Title": "HSE RCA DOCX Generator",
            "anthropic-version": "2023-06-01"  # Prompt caching için gerekli
        }

        # Anthropic Prompt Caching - sistem promptu cache'le (maliyeti %90 düşürür)
        base_payload = {
            "model": self.model,
            "messages": [
                {
                    "role": "system",
                    "content": [
                        {
                            "type": "text",
                            "text": CONTENT_SYSTEM_PROMPT,
                            "cache_control": {"type": "ephemeral"}  # Bu promptu 5 dakika cache'le
                        }
                    ]
                },
                {"role": "user", "content": user_msg}
            ],
            "max_tokens": 32000,
            "temperature": 0.3,
            "stream": False  # Non-streaming daha hızlı ve güvenilir
        }
        # Geçerli JSON zorunlu: kaçmayan tırnak / virgül hatalarını azaltır (OpenRouter/OpenAI uyumlu)
        use_json_object = os.getenv("OPENROUTER_DOCX_JSON_OBJECT", "1").strip() not in (
            "0",
            "false",
            "no",
        )
        if use_json_object:
            base_payload["response_format"] = {"type": "json_object"}

        print("-" * 50)

        deterministic_fallback = self._build_deterministic_fallback_content(raw_data, lang)
        def _request_and_parse(payload: Dict, label: str) -> Optional[Dict]:
            try:
                response = requests.post(
                    self.api_url,
                    headers=headers,
                    json=payload,
                    timeout=600
                )

                # Bazı modeller response_format ile 400 döndürebilir — bir kez düz istekle yeniden dene
                if response.status_code == 400 and use_json_object and "response_format" in payload:
                    detail = (response.text or "")[:500]
                    print(
                        f"\n  ⚠️  [{label}] response_format(json_object) reddedildi (400), "
                        f"format olmadan tekrar deneniyor...\n  Sunucu: {detail}"
                    )
                    retry_payload = {k: v for k, v in payload.items() if k != "response_format"}
                    response = requests.post(
                        self.api_url,
                        headers=headers,
                        json=retry_payload,
                        timeout=600
                    )

                if not response.ok:
                    snippet = (response.text or "")[:1200].replace("\n", " ")
                    print(f"\n OpenRouter HTTP {response.status_code} — URL: {self.api_url} [{label}]")
                    print(f" Yanıt özeti: {snippet}")
                    print("-" * 50)
                    return None

                try:
                    result = response.json()
                except ValueError:
                    snippet = (response.text or "")[:1200].replace("\n", " ")
                    print("\n OpenRouter yanıtı JSON değil (HTML veya hata sayfası olabilir).")
                    print(f" URL: {self.api_url}")
                    print(f" [{label}] Önizleme: {snippet}")
                    print("-" * 50)
                    return None

                try:
                    from shared.usage_context import record_openrouter_json

                    record_openrouter_json(
                        result,
                        reason="report_html",
                        model=self.model,
                        operation_label=f"Rapor LLM ({label})",
                    )
                except Exception:  # noqa: BLE001
                    pass

                err_obj = result.get("error")
                if err_obj:
                    print(f"\n OpenRouter API error [{label}]: {err_obj}")
                    print("-" * 50)
                    return None

                if "choices" in result and len(result["choices"]) > 0:
                    full_text = result["choices"][0].get("message", {}).get("content", "") or ""

                    if not full_text.strip():
                        print(f"\n OpenRouter: choices[0].message.content boş [{label}].")
                        print("-" * 50)
                        return None

                    # İçeriği ekrana yazdır (debug için)
                    print(full_text[:500] + "..." if len(full_text) > 500 else full_text)
                    print(f"\n Toplam karakter: {len(full_text)}")
                    print("-" * 50)

                    return self._parse_json_response(full_text)

                print(f"\n Geçersiz API yanıtı (choices yok) [{label}]: {str(result)[:800]}")
                print("-" * 50)
                return None

            except requests.exceptions.RequestException as e:
                print(f"\n OpenRouter API hatası [{label}]: {e}")
                print("-" * 50)
                return None

        # 1) İlk deneme
        parsed = _request_and_parse(base_payload, "attempt-1")
        if parsed:
            merged = self._merge_with_fallback_content(parsed, deterministic_fallback)
            pinned = self._pin_agent_why_chains(merged, raw_data)
            return self._enforce_taxonomy_titles(pinned, raw_data)

        # 2) Parse/format bozuksa, daha katı prompt ve düşük temperature ile bir kez yeniden dene
        retry_payload = copy.deepcopy(base_payload)
        retry_payload["temperature"] = 0.0
        retry_payload["messages"] = list(base_payload.get("messages", []))
        retry_user_msg = (
            user_msg
            + "\n\nÖNEMLİ: Önceki yanıt geçersiz JSON idi. "
              "Bu kez SADECE tek bir geçerli JSON nesnesi döndür. "
              "Eksik virgül, kapanmayan tırnak, trailing comma, markdown/code fence ekleme."
        )
        if len(retry_payload["messages"]) >= 2:
            retry_payload["messages"][1] = {"role": "user", "content": retry_user_msg}

        print("  🔁 İlk deneme parse edilemedi, katı JSON modunda ikinci deneme yapılıyor...")
        parsed_retry = _request_and_parse(retry_payload, "attempt-2")
        if parsed_retry:
            merged = self._merge_with_fallback_content(parsed_retry, deterministic_fallback)
            pinned = self._pin_agent_why_chains(merged, raw_data)
            return self._enforce_taxonomy_titles(pinned, raw_data)

        print("  ⚠️  Tüm denemeler başarısız, deterministic fallback rapora düşülüyor (boş rapor engellendi).")
        pinned = self._pin_agent_why_chains(deterministic_fallback, raw_data)
        return self._enforce_taxonomy_titles(pinned, raw_data)

    def _parse_json_response(self, text: str) -> Optional[Dict]:
        last_error: Optional[Exception] = None
        stripped = (text or "").strip()

        # 1) İlk markdown code-fence (non-greedy; birden fazla ``` varsa güvenli)
        m = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", stripped)
        if m:
            inner = m.group(1).strip()
            try:
                return json.loads(inner)
            except json.JSONDecodeError as e:
                last_error = e
            m_greedy = re.search(r"```(?:json)?\s*([\s\S]+)\s*```", stripped)
            if m_greedy and m_greedy.group(1) != m.group(1):
                try:
                    return json.loads(m_greedy.group(1).strip())
                except json.JSONDecodeError as e2:
                    last_error = e2

        try:
            return json.loads(stripped)
        except json.JSONDecodeError as e:
            last_error = e

        # İlk geçerli JSON nesnesini ayıkla (sonrasında model ek metin bırakmış olabilir)
        start_obj = stripped.find("{")
        if start_obj != -1:
            try:
                obj, _ = json.JSONDecoder().raw_decode(stripped[start_obj:])
                if isinstance(obj, dict):
                    return obj
            except json.JSONDecodeError as e:
                last_error = e

        start = stripped.find("{")
        end = stripped.rfind("}")
        if start != -1 and end != -1 and end > start:
            candidate = stripped[start:end + 1]
            try:
                return json.loads(candidate)
            except json.JSONDecodeError as e:
                last_error = e
                cleaned = re.sub(r",\s*([}\]])", r"\1", candidate)
                try:
                    return json.loads(cleaned)
                except json.JSONDecodeError as e2:
                    last_error = e2

        fallback = extract_json_from_response(stripped, default={})
        if fallback:
            return fallback

        err_msg = f"{type(last_error).__name__}: {last_error}" if last_error else "unknown"
        print(f"  ❌ JSON parse hatası: {err_msg}")
        if last_error is not None and hasattr(last_error, "lineno") and hasattr(last_error, "colno"):
            print(f"     konum: satır {last_error.lineno}, sütun {last_error.colno}")
        head = stripped[:300].replace("\n", " ")
        tail = stripped[-300:].replace("\n", " ") if len(stripped) > 300 else ""
        print(f"     baş: {head}")
        if tail:
            print(f"     son: {tail}")
        return None

    def _build_docx(self, content: Dict, output_path: str, lang: Optional[Dict] = None) -> None:
        lang = lang or {"code": "tr", "name": "Turkish", "rtl": False}
        lang_code = (lang.get("code") or "tr").lower()
        set_report_lang(lang_code)
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

        _build_cover(doc, content.get("cover", {}), lang_code)
        _build_executive_summary(doc, content.get("executive_summary", {}), content.get("root_causes", []))
        _build_incident_details(doc, content.get("incident_details", {}))
        _build_analysis_method(doc, content.get("analysis_method", {}))
        branches = content.get("branches", [])
        if branches:
            _build_branches(doc, branches)
        
        # Meta root cause (varsa)
        meta_root = content.get("meta_root_cause", {})
        if meta_root and meta_root.get("exists"):
            _build_meta_root_cause(doc, meta_root)
        
        _build_contributing_factors(doc, content.get("contributing_factors", []))
        _build_corrective_actions(doc, content.get("corrective_actions", []))
        _build_lessons_learned(doc, content.get("lessons_learned", {}))
        _build_conclusion(doc, content.get("conclusion", {}))
        _build_signature_page(doc)
        _translate_docx_static_labels(doc, lang_code)

        # RTL dil desteği (Arapça, İbranice vb.)
        if lang.get("rtl"):
            try:
                from lxml import etree as _etree  # python-docx already depends on lxml
                WNAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                for para in doc.paragraphs:
                    pPr = para._p.get_or_add_pPr()
                    bidi_tag = f"{{{WNAMESPACE}}}bidi"
                    if pPr.find(bidi_tag) is None:
                        _etree.SubElement(pPr, bidi_tag)
                    jc_tag = f"{{{WNAMESPACE}}}jc"
                    existing_jc = pPr.find(jc_tag)
                    if existing_jc is None:
                        jc = _etree.SubElement(pPr, jc_tag)
                        jc.set(f"{{{WNAMESPACE}}}val", "right")
            except ImportError:
                print("  lxml bulunamadı — RTL DOCX desteği atlandı")

        doc.save(output_path)
        print(f" Dosya kaydedildi: {output_path}")
    
    def _build_decision_tree(self, data: Dict, output_path: str) -> None:
        """5-Why Decision Tree HTML'ini oluşturur."""
        try:
            # RCA verisini çıkar
            rca_data = None
            incident_title = "Kaza Analizi"
            
            if "part3_rca" in data:
                rca_data = data["part3_rca"]
                # Olay başlığını part1'den al
                if "part1" in data and "overview" in data["part1"]:
                    overview = data["part1"]["overview"]
                    incident_title = overview.get("what_happened", "Kaza Analizi")
            elif "analysis_branches" in data:
                rca_data = data
                incident_title = rca_data.get("incident_event", "Kaza Analizi")
            
            if not rca_data:
                print("  Uyarı: RCA verisi bulunamadı, decision tree oluşturulamadı")
                return

            tree_payload = dict(rca_data) if isinstance(rca_data, dict) else {}
            part1 = data.get("part1") if isinstance(data.get("part1"), dict) else {}
            overview = part1.get("overview") if isinstance(part1.get("overview"), dict) else {}
            tree_payload["part1"] = part1
            try:
                from agents.report_text_sanitize import full_incident_narrative_for_tree
            except ImportError:
                from .report_text_sanitize import full_incident_narrative_for_tree
            narrative_candidates = [
                overview.get("what_happened"),
                part1.get("description"),
                tree_payload.get("incident_summary"),
                tree_payload.get("incident_event"),
            ]
            best_narrative = ""
            for src in narrative_candidates:
                if isinstance(src, str) and src.strip():
                    prepared = full_incident_narrative_for_tree(src.strip())
                    if len(prepared) > len(best_narrative):
                        best_narrative = prepared
            if best_narrative:
                tree_payload["incident_summary"] = best_narrative
            
            # Decision tree HTML'ini oluştur
            from agents.decision_tree_mermaid import DecisionTreeGenerator
            gen = DecisionTreeGenerator()
            gen.generate_html(
                rca_data=tree_payload,
                output_path=output_path,
                incident_title=incident_title
            )
            
        except Exception as e:
            print(f"  Uyarı: Decision tree oluşturulurken hata: {e}")

    def _prepare_content_for_display(self, content: Dict) -> Dict:
        """Kapak özeti ve olay detayları tablosunu sadeleştirir (yalnızca olay özeti)."""
        if not isinstance(content, dict):
            return content
        out = copy.deepcopy(content)
        cover = out.setdefault("cover", {})
        exec_sum = out.get("executive_summary") or {}
        raw_short = (
            cover.get("incident_summary_short")
            or exec_sum.get("what_happened")
            or ""
        )
        short = short_incident_summary(strip_hse_codes(str(raw_short)), 360)
        if not short:
            short = short_incident_summary(
                strip_hse_codes(str(exec_sum.get("what_happened") or "")), 360
            )
        cover["incident_summary_short"] = short

        details = out.setdefault("incident_details", {})
        event_table = dict(details.get("event_table") or {})
        for bulky_key in ("Özet", "Olay Özeti", "Summary", "Incident Summary"):
            val = event_table.pop(bulky_key, None)
            if val and not short:
                short = short_incident_summary(strip_hse_codes(str(val)), 360)
                cover["incident_summary_short"] = short
        details["event_table"] = event_table
        details["event_summary_only"] = short or details.get("event_summary_only") or ""
        return out

    def _build_html(self, content: Dict, output_path: str, lang: Optional[Dict] = None, investigation_data: Optional[Dict] = None) -> None:
        """Düzenlenebilir HTML rapor oluşturur."""
        content = self._prepare_content_for_display(content)
        html = self._generate_html_template(content, lang, investigation_data)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)

    def _html_branding_extras(
        self, investigation_data: Optional[Dict] = None
    ) -> tuple[str, str, str, str, str]:
        """Returns (watermark_css, watermark_html, cover_logo_html, cover_tpl_css, cover_class)."""
        layout = resolve_report_layout(investigation_data)
        wm_mode = str(layout.get("watermark_mode") or "final").lower()
        wm_label = ""
        if wm_mode == "draft":
            wm_label = "DRAFT"
        elif wm_mode == "final":
            wm_label = "FINAL"
        watermark_css = ""
        watermark_html = ""
        if wm_label:
            watermark_css = f"""
        body::before {{
            content: "{wm_label}";
            position: fixed;
            top: 42%;
            left: 50%;
            transform: translate(-50%, -50%) rotate(-32deg);
            font-size: 96px;
            font-weight: 800;
            color: rgba(44, 82, 130, 0.08);
            z-index: 0;
            pointer-events: none;
            letter-spacing: 0.2em;
        }}
        .container {{ position: relative; z-index: 1; }}
            """
            watermark_html = f'<div class="report-watermark" aria-hidden="true">{wm_label}</div>'
        logo_url = (layout.get("logo_url") or "").strip()
        cover_logo_html = ""
        if logo_url:
            cover_logo_html = (
                f'<div class="tenant-logo" style="margin-bottom:20px;">'
                f'<img src="{logo_url}" alt="Logo" style="max-height:72px;max-width:240px;" />'
                f"</div>"
            )
        cover_tpl = str(layout.get("cover_template") or "standard").lower()
        if cover_tpl not in ("standard", "formal", "executive", "minimal"):
            cover_tpl = "standard"
        cover_class = f"cover cover--{cover_tpl}"
        cover_tpl_css = f"""
        .cover--formal {{
            background: linear-gradient(160deg, #0f172a 0%, #1e293b 55%, #334155 100%);
            border-bottom: 6px solid #cbd5e1;
        }}
        .cover--formal h1 {{ letter-spacing: 0.08em; font-family: Georgia, 'Times New Roman', serif; }}
        .cover--executive {{
            background: linear-gradient(180deg, #fff 0%, #f8fafc 100%);
            color: #0f172a;
            border-top: 8px solid #b45309;
            border-bottom: 1px solid #e2e8f0;
        }}
        .cover--executive .confidential-banner {{ background: #b45309; color: #fff; }}
        .cover--executive .info-item {{ background: #fff; border-left-color: #b45309; color: #334155; }}
        .cover--executive .info-label {{ color: #b45309; }}
        .cover--minimal {{
            background: #f8fafc;
            color: #1e293b;
            border-bottom: 2px solid #e2e8f0;
            padding-top: 40px;
        }}
        .cover--minimal .confidential-banner {{
            background: transparent;
            color: #64748b;
            border: 1px dashed #cbd5e0;
        }}
        .cover--minimal h1 {{ text-transform: none; font-size: 2.1em; color: #0f172a; }}
        """
        return watermark_css, watermark_html, cover_logo_html, cover_tpl_css, cover_class

    def _generate_html_template(self, content: Dict, lang: Optional[Dict] = None, investigation_data: Optional[Dict] = None) -> str:
        """Modern, responsive ve düzenlenebilir HTML rapor şablonu."""
        lang = lang or {"code": "tr", "name": "Turkish", "rtl": False, "html_lang": "tr"}
        html_lang = lang.get("html_lang", "tr")
        lang_code = (lang.get("code") or "tr").lower()
        set_report_lang(lang_code)
        is_rtl = lang.get("rtl", False)
        dir_attr = ' dir="rtl"' if is_rtl else ''
        watermark_css, watermark_html, cover_logo_html, cover_tpl_css, cover_class = self._html_branding_extras(investigation_data)
        rtl_css = """
        body { direction: rtl; text-align: right; }
        .container { direction: rtl; }
        .cover { direction: rtl; }
        th, td { text-align: right; }
        .meta-table td:first-child { font-weight: bold; }
        """ if is_rtl else ""

        cover = content.get("cover", {})
        executive_summary = content.get("executive_summary", {})
        incident_details = content.get("incident_details", {})
        analysis_method = content.get("analysis_method", {})
        branches = content.get("branches", [])
        root_causes = content.get("root_causes", [])
        root_causes = self._ensure_root_causes_from_branches(root_causes, branches)
        contributing_factors = content.get("contributing_factors", [])
        corrective_actions = content.get("corrective_actions", [])
        lessons_learned = content.get("lessons_learned", {})
        conclusion = content.get("conclusion", {})

        # HTML oluştur
        html = f"""<!DOCTYPE html>
<html lang="{html_lang}"{dir_attr}>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{cover.get('title', _label(lang_code, 'cover_title'))}</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #444;
            background: #f5f5f5;
            padding: 20px;
        }}
        {rtl_css}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
        }}
        
        .cover {{
            background: #2c5282;
            color: #f7fafc;
            padding: 56px 32px;
            text-align: center;
        }}
        
        .cover h1 {{
            font-size: 2.5em;
            margin-bottom: 20px;
            text-transform: uppercase;
        }}
        
        .cover .subtitle {{
            font-size: 1.2em;
            font-style: italic;
            margin-bottom: 30px;
        }}
        {cover_tpl_css}
        
        .confidential-banner {{
            background: #C0392B;
            color: white;
            padding: 15px;
            margin: 30px 0;
            font-weight: bold;
            text-align: center;
            border-radius: 5px;
        }}
        
        .info-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 15px;
            margin: 30px 0;
        }}
        
        .info-item {{
            background: #edf2f7;
            padding: 15px;
            border-left: 4px solid #2c5282;
        }}
        
        .info-label {{
            font-weight: bold;
            color: #2c5282;
            font-size: 0.9em;
        }}
        
        .info-value {{
            margin-top: 5px;
            color: #444;
        }}
        
        .incident-summary {{
            background: #edf2f7;
            color: #1a202c;
            padding: 18px 22px;
            margin: 24px auto;
            max-width: 720px;
            border-radius: 6px;
            border: 1px solid #cbd5e0;
            text-align: left;
        }}
        
        .incident-summary h3 {{
            color: #2c5282;
            margin: 0 0 10px 0;
            font-size: 1.05em;
        }}
        
        .incident-summary p {{
            margin: 0;
            line-height: 1.55;
            font-size: 0.95em;
        }}
        
        .event-summary-compact {{
            max-width: 720px;
            margin: 16px 0 24px 0;
            padding: 14px 18px;
            background: #f7fafc;
            border: 1px solid #e2e8f0;
            border-radius: 6px;
            line-height: 1.55;
            font-size: 0.95em;
        }}
        
        .event-summary-compact .label {{
            font-weight: bold;
            color: #2c5282;
            display: block;
            margin-bottom: 8px;
        }}
        
        .content {{
            padding: 40px;
        }}
        
        .section {{
            margin: 40px 0;
            page-break-inside: avoid;
        }}
        
        .section-header {{
            background: #2c5282;
            color: #f7fafc;
            padding: 15px 20px;
            margin: 30px 0 20px 0;
            font-size: 1.3em;
            font-weight: bold;
            text-transform: uppercase;
        }}
        
        .subsection-header {{
            color: #2c5282;
            font-size: 1.1em;
            font-weight: bold;
            margin: 25px 0 15px 0;
            padding-bottom: 5px;
            border-bottom: 2px solid #cbd5e0;
        }}
        
        .paragraph {{
            margin: 15px 0;
            text-align: justify;
            line-height: 1.8;
        }}
        
        .colored-box {{
            margin: 20px 0;
            border-radius: 5px;
            overflow: hidden;
        }}
        
        .box-header {{
            padding: 15px;
            font-weight: bold;
            color: white;
        }}
        
        .box-content {{
            background: #F5F5F5;
            padding: 20px;
            line-height: 1.75;
        }}
        .box-content strong {{
            color: #1B3A5C;
            font-weight: 700;
        }}
        .box-content .rc-point {{
            margin: 0.65em 0;
            padding: 0.35em 0 0.35em 0.85em;
            border-left: 3px solid #2E6DA4;
        }}
        .box-content .report-para {{
            margin: 0.5em 0;
        }}
        
        .box-red .box-header {{ background: #C0392B; }}
        .box-orange .box-header {{ background: #E67E22; }}
        .box-green .box-header {{ background: #27AE60; }}
        .box-blue .box-header {{ background: #2E6DA4; }}
        
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        
        th {{
            background: #1B3A5C;
            color: white;
            padding: 12px;
            text-align: left;
            font-weight: bold;
        }}
        
        td {{
            padding: 12px;
            border-bottom: 1px solid #ddd;
        }}
        
        tr:nth-child(even) {{
            background: #F5F5F5;
        }}
        
        tr:hover {{
            background: #E8F4F8;
        }}
        
        .timeline {{
            margin: 20px 0;
        }}
        
        .timeline-item {{
            display: flex;
            margin: 15px 0;
            padding: 15px;
            background: white;
            border-left: 4px solid #2E6DA4;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        
        .timeline-time {{
            font-weight: bold;
            color: #2E6DA4;
            min-width: 80px;
            font-size: 1.1em;
        }}
        
        .timeline-event {{
            flex: 1;
            margin-left: 20px;
        }}
        
        .why-chain {{
            margin: 20px 0;
        }}
        
        .why-item {{
            margin: 15px 0;
            padding: 15px;
            border-left: 4px solid #E67E22;
            background: #FFF8F0;
        }}
        
        .why-number {{
            font-weight: bold;
            color: #E67E22;
            font-size: 1.1em;
        }}
        
        .why-question {{
            font-weight: bold;
            margin: 5px 0;
            color: #444;
        }}
        
        .why-answer {{
            margin: 5px 0;
            padding-left: 20px;
        }}
        
        .why-code {{
            display: inline-block;
            background: #E67E22;
            color: white;
            padding: 3px 10px;
            border-radius: 3px;
            font-size: 0.85em;
            margin-top: 5px;
        }}
        
        .root-cause-box {{
            margin: 30px 0;
            border-radius: 5px;
            overflow: hidden;
            box-shadow: 0 4px 10px rgba(0,0,0,0.15);
        }}
        
        .root-cause-header {{
            padding: 20px;
            color: white;
            font-size: 1.2em;
            font-weight: bold;
        }}
        
        .root-cause-content {{
            background: white;
            padding: 25px;
        }}
        
        .root-cause-1 .root-cause-header {{ background: #C0392B; }}
        .root-cause-2 .root-cause-header {{ background: #E67E22; }}
        .root-cause-3 .root-cause-header {{ background: #27AE60; }}
        .root-cause-4 .root-cause-header {{ background: #2E6DA4; }}
        
        ul.bullet-list {{
            margin: 15px 0;
            padding-left: 30px;
        }}
        
        ul.bullet-list li {{
            margin: 8px 0;
            line-height: 1.6;
        }}
        
        .priority-urgent {{
            background: #C0392B;
            color: white;
            padding: 5px 10px;
            border-radius: 3px;
            font-weight: bold;
            font-size: 0.85em;
        }}
        
        .priority-high {{
            background: #E67E22;
            color: white;
            padding: 5px 10px;
            border-radius: 3px;
            font-weight: bold;
            font-size: 0.85em;
        }}
        
        .priority-medium {{
            background: #27AE60;
            color: white;
            padding: 5px 10px;
            border-radius: 3px;
            font-weight: bold;
            font-size: 0.85em;
        }}
        
        .priority-low {{
            background: #2E6DA4;
            color: white;
            padding: 5px 10px;
            border-radius: 3px;
            font-weight: bold;
            font-size: 0.85em;
        }}
        
        .impact-high {{ color: #C0392B; font-weight: bold; }}
        .impact-medium {{ color: #E67E22; font-weight: bold; }}
        .impact-low {{ color: #27AE60; font-weight: bold; }}
        
        .signature-section {{
            margin: 40px 0;
        }}
        
        .signature-table td {{
            padding: 30px 15px;
        }}
        
        .signature-line {{
            border-top: 2px solid #444;
            margin-top: 60px;
            padding-top: 10px;
            text-align: center;
        }}

        /*  OLAY FOTOĞRAFLARI  */
        .photo-section {{
            margin: 40px 0;
            page-break-before: always;
        }}
        .photo-page {{
            page-break-after: always;
        }}
        .photo-page:last-child {{
            page-break-after: auto;
        }}
        .photo-grid {{
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 24px;
            margin-top: 20px;
        }}
        .photo-slot {{
            border: 2px dashed #2E6DA4;
            border-radius: 10px;
            min-height: 260px;
            display: flex;
            flex-direction: column;
            align-items: center;
            justify-content: center;
            background: #F8FBFF;
            cursor: pointer;
            transition: border-color 0.2s, background 0.2s;
            overflow: hidden;
            position: relative;
        }}
        .photo-slot:hover {{
            border-color: #1B3A5C;
            background: #EBF3FB;
        }}
        .photo-slot.has-photo {{
            border-style: solid;
            border-color: #1B3A5C;
            background: #fff;
        }}
        .photo-slot img {{
            width: 100%;
            height: 240px;
            object-fit: cover;
            border-radius: 8px 8px 0 0;
            display: block;
        }}
        .photo-slot .upload-hint {{
            display: flex;
            flex-direction: column;
            align-items: center;
            gap: 10px;
            color: #2E6DA4;
            pointer-events: none;
        }}
        .photo-slot .upload-hint .icon {{
            font-size: 3em;
        }}
        .photo-slot .upload-hint .label {{
            font-size: 0.95em;
            font-weight: 600;
        }}
        .photo-slot .upload-hint .sub {{
            font-size: 0.78em;
            color: #888;
        }}
        .photo-caption {{
            width: 100%;
            padding: 8px 10px;
            border: none;
            border-top: 1px solid #ddd;
            font-size: 0.85em;
            background: #fff;
            color: #333;
            text-align: center;
            outline: none;
            border-radius: 0 0 8px 8px;
        }}
        .photo-caption::placeholder {{
            color: #aaa;
        }}
        .photo-remove {{
            position: absolute;
            top: 6px;
            right: 8px;
            background: rgba(192,57,43,0.85);
            color: white;
            border: none;
            border-radius: 50%;
            width: 26px;
            height: 26px;
            font-size: 0.9em;
            cursor: pointer;
            display: none;
            align-items: center;
            justify-content: center;
            z-index: 5;
        }}
        .photo-slot.has-photo .photo-remove {{
            display: flex;
        }}
        .photo-page-title {{
            font-size: 1em;
            color: #888;
            margin-bottom: 4px;
        }}
        @media print {{
            .photo-slot .upload-hint {{ display: none; }}
            .photo-remove {{ display: none !important; }}
            .photo-caption {{ border-top: 1px solid #ccc; }}
        }}
        
        .comparison-table td:first-child {{
            background: #D6E4F0 !important;
            color: #1B3A5C;
            font-weight: bold;
        }}
        
        .comparison-table .current {{
            background: #FFE6E6 !important;
            color: #C0392B;
        }}
        
        .comparison-table .target {{
            background: #E8F8F0 !important;
            color: #27AE60;
            font-weight: bold;
        }}
        
        /* Navigasyon Menüsü */
        .nav-menu {{
            position: fixed;
            top: 20px;
            right: 20px;
            background: white;
            padding: 15px;
            border-radius: 8px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.2);
            z-index: 1000;
            max-width: 250px;
        }}
        
        .nav-menu h3 {{
            margin: 0 0 10px 0;
            font-size: 1em;
            color: #1B3A5C;
            border-bottom: 2px solid #1B3A5C;
            padding-bottom: 5px;
        }}
        
        .nav-menu ul {{
            list-style: none;
            padding: 0;
            margin: 0;
        }}
        
        .nav-menu li {{
            margin: 8px 0;
        }}
        
        .nav-menu a {{
            color: #2E6DA4;
            text-decoration: none;
            font-size: 0.9em;
            display: block;
            padding: 5px;
            border-radius: 3px;
            transition: all 0.2s;
        }}
        
        .nav-menu a:hover {{
            background: #D6E4F0;
            padding-left: 10px;
        }}
        
        .nav-toggle {{
            position: fixed;
            top: 20px;
            right: 20px;
            background: #1B3A5C;
            color: white;
            border: none;
            padding: 12px 20px;
            border-radius: 5px;
            cursor: pointer;
            font-size: 0.9em;
            z-index: 999;
            box-shadow: 0 2px 10px rgba(0,0,0,0.2);
        }}
        
        .nav-toggle:hover {{
            background: #2E6DA4;
        }}
        
        /* Düzenleme Toolbar */
        .edit-toolbar {{
            position: fixed;
            bottom: 20px;
            left: 50%;
            transform: translateX(-50%);
            background: white;
            padding: 15px;
            border-radius: 8px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.2);
            z-index: 1000;
            display: none;
        }}
        
        .edit-toolbar.active {{
            display: flex;
            gap: 10px;
            align-items: center;
        }}
        
        .toolbar-btn {{
            padding: 8px 15px;
            border: none;
            border-radius: 5px;
            cursor: pointer;
            font-size: 0.9em;
            transition: all 0.2s;
        }}
        
        .toolbar-btn:hover {{
            transform: translateY(-2px);
            box-shadow: 0 2px 5px rgba(0,0,0,0.2);
        }}
        
        .btn-save {{
            background: #27AE60;
            color: white;
        }}
        
        .btn-print {{
            background: #2E6DA4;
            color: white;
        }}
        
        .btn-export {{
            background: #E67E22;
            color: white;
        }}
        
        .btn-reset {{
            background: #C0392B;
            color: white;
        }}
        
        .btn-edit-mode {{
            background: #9B59B6;
            color: white;
        }}
        
        /* Sayfa Numaraları (Yazdırma için) */
        @page {{
            margin: 2cm;
            @bottom-right {{
                content: "Sayfa " counter(page) " / " counter(pages);
                font-size: 10pt;
                color: #666;
            }}
            @bottom-left {{
                content: "HSE Kök Neden Analizi - {cover.get('ref_no', 'N/A')}";
                font-size: 10pt;
                color: #666;
            }}
        }}
        
        /* Yazdırma Ayarları */
        @media print {{
            body {{
                background: white;
                padding: 0;
            }}
            
            .container {{
                box-shadow: none;
                max-width: none;
            }}
            
            .section {{
                page-break-inside: avoid;
            }}
            
            .section-header {{
                page-break-after: avoid;
            }}
            
            .root-cause-box {{
                page-break-inside: avoid;
            }}
            
            .nav-menu, .nav-toggle, .edit-toolbar {{
                display: none !important;
            }}
            
            /* Sayfa numaraları için footer */
            .page-footer {{
                position: fixed;
                bottom: 0;
                left: 0;
                right: 0;
                text-align: center;
                font-size: 10pt;
                color: #666;
                padding: 10px;
                border-top: 1px solid #ddd;
            }}
            
            /* Bölüm başlarında sayfa ayırıcı */
            .section-header {{
                page-break-before: always;
            }}
            
            .cover {{
                page-break-after: always;
            }}
        }}
        
        /* Düzenlenebilir alanlar için */
        [contenteditable="true"] {{
            outline: none;
            transition: background 0.2s;
            position: relative;
        }}
        
        [contenteditable="true"]:hover {{
            background: #FFFACD;
        }}
        
        [contenteditable="true"]:focus {{
            background: #FFFFE0;
            border: 1px dashed #E67E22;
            padding: 5px;
        }}
        
        [contenteditable="true"]:hover::after {{
            content: " Düzenlemek için tıklayın";
            position: absolute;
            top: -25px;
            left: 0;
            background: #E67E22;
            color: white;
            padding: 3px 8px;
            border-radius: 3px;
            font-size: 0.75em;
            white-space: nowrap;
            z-index: 100;
        }}
        
        .edit-hint {{
            color: #999;
            font-size: 0.85em;
            font-style: italic;
            margin-top: 5px;
        }}
        
        /* Scroll-to-top button */
        .scroll-top {{
            position: fixed;
            bottom: 80px;
            right: 20px;
            background: #1B3A5C;
            color: white;
            border: none;
            width: 50px;
            height: 50px;
            border-radius: 50%;
            cursor: pointer;
            font-size: 1.5em;
            display: none;
            box-shadow: 0 2px 10px rgba(0,0,0,0.2);
            z-index: 998;
        }}
        
        .scroll-top:hover {{
            background: #2E6DA4;
        }}
        
        .scroll-top.visible {{
            display: block;
        }}
        
        /* Highlight effect for navigation */
        .section.highlighted {{
            animation: highlight 1s ease-in-out;
        }}
        
        @keyframes highlight {{
            0% {{ background: transparent; }}
            50% {{ background: #FFFACD; }}
            100% {{ background: transparent; }}
        }}
        {watermark_css}
    </style>
</head>
<body>
    {watermark_html}
    <!-- Navigasyon Toggle Butonu -->
    <button class="nav-toggle" onclick="toggleNav()"> {_label(lang_code, 'nav_toc')}</button>
    
    <!-- Navigasyon Menüsü -->
    <div class="nav-menu" id="navMenu" style="display: none;">
        <h3>{_label(lang_code, 'nav_toc_title')}</h3>
        <ul>
            <li><a href="#cover" onclick="scrollToSection('cover')"> {_label(lang_code, 'nav_cover')}</a></li>
            <li><a href="#executive-summary" onclick="scrollToSection('executive-summary')"> {_label(lang_code, 'nav_executive_summary')}</a></li>
            <li><a href="#incident-details" onclick="scrollToSection('incident-details')"> {_label(lang_code, 'nav_incident_details')}</a></li>
            <li><a href="#analysis-method" onclick="scrollToSection('analysis-method')"> {_label(lang_code, 'nav_analysis_method')}</a></li>
            <li><a href="#branches" onclick="scrollToSection('branches')"> {_label(lang_code, 'nav_branches')}</a></li>
            <li><a href="#contributing-factors" onclick="scrollToSection('contributing-factors')"> {_label(lang_code, 'nav_contributing')}</a></li>
            <li><a href="#corrective-actions" onclick="scrollToSection('corrective-actions')"> {_label(lang_code, 'nav_corrective')}</a></li>
            <li><a href="#lessons-learned" onclick="scrollToSection('lessons-learned')"> {_label(lang_code, 'nav_lessons')}</a></li>
            <li><a href="#conclusion" onclick="scrollToSection('conclusion')"> {_label(lang_code, 'nav_conclusion')}</a></li>
            <li><a href="#signatures" onclick="scrollToSection('signatures')"> {_label(lang_code, 'nav_signatures')}</a></li>
            <li><a href="#incident-photos" onclick="scrollToSection('incident-photos')"> {_label(lang_code, 'nav_photos')}</a></li>
        </ul>
    </div>
    
    <!-- Düzenleme Toolbar -->
    <div class="edit-toolbar" id="editToolbar">
        <button class="toolbar-btn btn-edit-mode" onclick="toggleEditMode()">
            <span id="editModeText"> {_label(lang_code, 'toolbar_edit_off')}</span>
        </button>
        <button class="toolbar-btn btn-save" onclick="saveReport()" title="{_label(lang_code, 'toolbar_save')}">
             {_label(lang_code, 'toolbar_save')}
        </button>
        <button class="toolbar-btn btn-export" onclick="exportWord()" title="{_label(lang_code, 'toolbar_word')}">
             {_label(lang_code, 'toolbar_word')}
        </button>
        <button class="toolbar-btn btn-export" onclick="exportHTML()" title="{_label(lang_code, 'toolbar_html')}">
             {_label(lang_code, 'toolbar_html')}
        </button>
        <button class="toolbar-btn btn-reset" onclick="resetReport()" title="{_label(lang_code, 'toolbar_reset')}">
             {_label(lang_code, 'toolbar_reset')}
        </button>
    </div>
    
    <!-- Scroll to Top Button -->
    <button class="scroll-top" id="scrollTopBtn" onclick="scrollToTop()">↑</button>
    
    <div class="container">
        <!-- KAPAK SAYFASI -->
        <div class="{cover_class}" id="cover">
            {cover_logo_html}
            <h1 contenteditable="true">{cover.get('title', _label(lang_code, 'cover_title'))}</h1>
            <div class="subtitle" contenteditable="true">{cover.get('subtitle', _label(lang_code, 'cover_subtitle'))}</div>
            
            <div class="confidential-banner" contenteditable="true">
                {cover.get('confidentiality', _label(lang_code, 'cover_confidentiality'))}
            </div>
            
            <div class="info-grid">
                <div class="info-item">
                    <div class="info-label">{_label(lang_code, 'ref_no')}</div>
                    <div class="info-value" contenteditable="true">{cover.get('ref_no', 'N/A')}</div>
                </div>
                <div class="info-item">
                    <div class="info-label">{_label(lang_code, 'date')}</div>
                    <div class="info-value" contenteditable="true">{cover.get('date', 'N/A')}</div>
                </div>
                <div class="info-item">
                    <div class="info-label">{_label(lang_code, 'location')}</div>
                    <div class="info-value" contenteditable="true">{cover.get('location', 'N/A')}</div>
                </div>
                <div class="info-item">
                    <div class="info-label">{_label(lang_code, 'incident_type')}</div>
                    <div class="info-value" contenteditable="true">{cover.get('incident_type', 'N/A')}</div>
                </div>
            </div>
            
            <div class="incident-summary">
                <h3>{_label(lang_code, 'incident_summary')}</h3>
                <p contenteditable="true">{cover.get('incident_summary_short', '')}</p>
            </div>
        </div>
        
        <!-- İÇERİK -->
        <div class="content">
"""

        # 1. YÖNETİCİ ÖZETİ
        html += self._html_executive_summary(executive_summary, root_causes)
        
        # 2. OLAY BİLGİLERİ
        html += self._html_incident_details(incident_details)
        
        # 3. ANALİZ YÖNTEMİ
        html += self._html_analysis_method(analysis_method)
        
        # 4-N. DALLAR
        html += self._html_branches(branches)
        
        # META KÖK NEDEN KALDIRILDI - İstenmeyen karmaşıklık
        
        # Branch sayısına göre dinamik bölüm numaraları
        # 1: Executive, 2: Incident, 3: Method, 4..: Branches
        next_section_no = 4 + len(branches)

        # N+1. SİSTEMSEL FAKTÖRLER
        html += self._html_contributing_factors(contributing_factors, section_no=next_section_no)
        next_section_no += 1
        
        # N+2. DÜZELTİCİ FAALİYETLER
        html += self._html_corrective_actions(corrective_actions, section_no=next_section_no)
        next_section_no += 1
        
        # N+3. ÇIKARILAN DERSLER
        html += self._html_lessons_learned(lessons_learned, section_no=next_section_no)
        next_section_no += 1
        
        # N+4. SONUÇ
        html += self._html_conclusion(conclusion, section_no=next_section_no)
        next_section_no += 1
        
        # N+5. İMZA SAYFASI
        html += self._html_signatures(section_no=next_section_no)
        next_section_no += 1

        # N+6. OLAY FOTOĞRAFLARI (2 sayfa × 4 foto)
        html += self._html_incident_photos()

        html += """
        </div>
    </div>
    
    <script>
        // Navigasyon toggle
        function toggleNav() {
            const menu = document.getElementById('navMenu');
            menu.style.display = menu.style.display === 'none' ? 'block' : 'none';
        }
        
        // Bölüme kaydır ve highlight
        function scrollToSection(sectionId) {
            const element = document.getElementById(sectionId);
            if (element) {
                element.scrollIntoView({ behavior: 'smooth', block: 'start' });
                element.classList.add('highlighted');
                setTimeout(() => element.classList.remove('highlighted'), 1000);
            }
            // Mobilde menüyü kapat
            if (window.innerWidth < 768) {
                document.getElementById('navMenu').style.display = 'none';
            }
        }
        
        // Scroll to top
        function scrollToTop() {
            window.scrollTo({ top: 0, behavior: 'smooth' });
        }
        
        // Scroll position tracking
        window.addEventListener('scroll', function() {
            const scrollBtn = document.getElementById('scrollTopBtn');
            if (window.pageYOffset > 300) {
                scrollBtn.classList.add('visible');
            } else {
                scrollBtn.classList.remove('visible');
            }
        });
        
        // Düzenleme modu toggle
        let editMode = false;
        const LBL_EDIT_ON = ' {_label(lang_code, "toolbar_edit_on")}';
        const LBL_EDIT_OFF = ' {_label(lang_code, "toolbar_edit_off")}';
        const LBL_NOTIFY_EDIT_ON = ' {_label(lang_code, "notify_edit_on")}';
        const LBL_NOTIFY_EDIT_OFF = ' {_label(lang_code, "notify_edit_off")}';
        const LBL_NOTIFY_SAVED = ' {_label(lang_code, "notify_saved")}';
        const LBL_NOTIFY_WORD_PREP = ' {_label(lang_code, "notify_word_prep")}';
        const LBL_NOTIFY_WORD_DL = ' {_label(lang_code, "notify_word_dl")}';
        const LBL_NOTIFY_HTML_DL = ' {_label(lang_code, "notify_html_dl")}';
        const LBL_COVER_TITLE = '{_label(lang_code, "cover_title")}';
        function toggleEditMode() {
            editMode = !editMode;
            const editableElements = document.querySelectorAll('[contenteditable]');
            const editModeText = document.getElementById('editModeText');
            const toolbar = document.getElementById('editToolbar');
            
            if (editMode) {
                editableElements.forEach(el => el.setAttribute('contenteditable', 'true'));
                editModeText.textContent = LBL_EDIT_ON;
                toolbar.classList.add('active');
                showNotification(LBL_NOTIFY_EDIT_ON, 'success');
            } else {
                editableElements.forEach(el => el.setAttribute('contenteditable', 'false'));
                editModeText.textContent = LBL_EDIT_OFF;
                toolbar.classList.remove('active');
                showNotification(LBL_NOTIFY_EDIT_OFF, 'info');
            }
        }
        
        // Raporu kaydet (localStorage)
        function saveReport() {
            const html = document.documentElement.outerHTML;
            const timestamp = new Date().toISOString();
            localStorage.setItem('hse_report_saved', html);
            localStorage.setItem('hse_report_saved_time', timestamp);
            showNotification(LBL_NOTIFY_SAVED, 'success');
            console.log('Rapor kaydedildi:', timestamp);
        }
        
        // Yazdır / PDF kaydet
        function printReport() {
            // Düzenleme modunu kapat
            if (editMode) {
                toggleEditMode();
            }
            
            showNotification(' Yazdırma ekranı açılıyor...', 'info');
            setTimeout(() => {
                window.print();
            }, 500);
        }
        
        // HTML olarak indir
        function exportHTML() {
            const html = document.documentElement.outerHTML;
            const blob = new Blob([html], { type: 'text/html' });
            const url = URL.createObjectURL(blob);
            const a = document.createElement('a');
            a.href = url;
            a.download = 'hse_report_' + new Date().getTime() + '.html';
            document.body.appendChild(a);
            a.click();
            document.body.removeChild(a);
            URL.revokeObjectURL(url);
            showNotification(LBL_NOTIFY_HTML_DL, 'success');
        }
        
        // PDF olarak indir (düzenlenebilir)
        function exportPDF() {
            // Düzenleme modunu kapat
            if (editMode) {
                toggleEditMode();
            }
            
            // html2pdf kütüphanesi yükle
            const script = document.createElement('script');
            script.src = 'https://cdnjs.cloudflare.com/ajax/libs/html2pdf.js/0.10.1/html2pdf.bundle.min.js';
            script.onload = function() {
                const element = document.querySelector('.container');
                const opt = {
                    margin: 10,
                    filename: 'hse_report_' + new Date().getTime() + '.pdf',
                    image: { type: 'jpeg', quality: 0.98 },
                    html2canvas: { scale: 2 },
                    jsPDF: { orientation: 'portrait', unit: 'mm', format: 'a4' }
                };
                html2pdf().set(opt).from(element).save();
                showNotification(' PDF dosyası indiriliyor...', 'success');
            };
            document.head.appendChild(script);
        }
        
        // Word olarak indir (DOCX)
        function exportWord() {
            // Düzenleme modunu kapat
            if (editMode) {
                toggleEditMode();
            }
            
            showNotification(LBL_NOTIFY_WORD_PREP, 'info');
            
            // docx kütüphanesi yükle
            const script = document.createElement('script');
            script.src = 'https://cdnjs.cloudflare.com/ajax/libs/docx/8.5.0/docx.min.js';
            script.onload = function() {
                try {
                    const html = document.querySelector('.container').innerHTML;
                    
                    // HTML'i temizle
                    const temp = document.createElement('div');
                    temp.innerHTML = html;
                    const text = temp.innerText;
                    
                    // Basit DOCX oluştur
                    const docx = new docx.Document({
                        sections: [{
                            properties: {},
                            children: [
                                new docx.Paragraph({
                                    text: LBL_COVER_TITLE,
                                    heading: docx.HeadingLevel.HEADING_1,
                                    alignment: docx.AlignmentType.CENTER,
                                    spacing: { after: 400 }
                                }),
                                new docx.Paragraph({
                                    text: text,
                                    spacing: { line: 360 }
                                })
                            ]
                        }]
                    });
                    
                    docx.Packer.toBlob(docx).then(blob => {
                        const url = URL.createObjectURL(blob);
                        const a = document.createElement('a');
                        a.href = url;
                        a.download = 'hse_report_' + new Date().getTime() + '.docx';
                        document.body.appendChild(a);
                        a.click();
                        document.body.removeChild(a);
                        URL.revokeObjectURL(url);
                        showNotification(LBL_NOTIFY_WORD_DL, 'success');
                    });
                } catch (error) {
                    showNotification(' Word export başarısız: ' + error.message, 'error');
                }
            };
            document.head.appendChild(script);
        }
        
        // Orijinal haline döndür
        function resetReport() {
            if (confirm(' Tüm değişiklikler kaybolacak. Orijinal rapora dönmek istediğinizden emin misiniz?')) {
                location.reload();
                showNotification(' Rapor sıfırlandı', 'info');
            }
        }
        
        // Bildirim göster
        function showNotification(message, type = 'info') {
            // Mevcut bildirimi kaldır
            const existing = document.querySelector('.notification');
            if (existing) {
                existing.remove();
            }
            
            // Yeni bildirim oluştur
            const notification = document.createElement('div');
            notification.className = 'notification notification-' + type;
            notification.textContent = message;
            notification.style.cssText = `
                position: fixed;
                top: 80px;
                right: 20px;
                background: ${type === 'success' ? '#27AE60' : type === 'error' ? '#C0392B' : '#2E6DA4'};
                color: white;
                padding: 15px 20px;
                border-radius: 5px;
                box-shadow: 0 4px 15px rgba(0,0,0,0.3);
                z-index: 10000;
                animation: slideIn 0.3s ease-out;
                font-weight: bold;
            `;
            
            document.body.appendChild(notification);
            
            // 3 saniye sonra kaldır
            setTimeout(() => {
                notification.style.animation = 'slideOut 0.3s ease-out';
                setTimeout(() => notification.remove(), 300);
            }, 3000);
        }
        
        // Animasyonlar
        const style = document.createElement('style');
        style.textContent = `
            @keyframes slideIn {
                from {
                    transform: translateX(400px);
                    opacity: 0;
                }
                to {
                    transform: translateX(0);
                    opacity: 1;
                }
            }
            @keyframes slideOut {
                from {
                    transform: translateX(0);
                    opacity: 1;
                }
                to {
                    transform: translateX(400px);
                    opacity: 0;
                }
            }
        `;
        document.head.appendChild(style);
        
        // Otomatik kaydetme
        let autoSaveTimeout;
        document.addEventListener('input', function(e) {
            if (e.target.hasAttribute('contenteditable')) {
                clearTimeout(autoSaveTimeout);
                autoSaveTimeout = setTimeout(() => {
                    const html = document.documentElement.outerHTML;
                    localStorage.setItem('hse_report_autosave', html);
                    localStorage.setItem('hse_report_autosave_time', new Date().toISOString());
                    console.log(' Otomatik kaydedildi:', new Date().toLocaleTimeString());
                }, 2000); // 2 saniye sonra otomatik kaydet
            }
        });
        
        // Sayfa yüklendiğinde toolbar'ı göster
        window.addEventListener('load', function() {
            document.getElementById('editToolbar').classList.add('active');
            
            // Kaydedilmiş rapor var mı kontrol et
            const savedTime = localStorage.getItem('hse_report_saved_time');
            if (savedTime) {
                console.log(' Son kayıt:', new Date(savedTime).toLocaleString('tr-TR'));
            }
            
            showNotification(' Rapor yüklendi - Düzenlemek için  butonuna tıklayın', 'info');
        });
        
        // Keyboard shortcuts
        document.addEventListener('keydown', function(e) {
            // Ctrl+S: Save
            if (e.ctrlKey && e.key === 's') {
                e.preventDefault();
                saveReport();
            }
            // Ctrl+E: Toggle edit mode
            if (e.ctrlKey && e.key === 'e') {
                e.preventDefault();
                toggleEditMode();
            }
            // Escape: Close nav
            if (e.key === 'Escape') {
                document.getElementById('navMenu').style.display = 'none';
            }
        });
        
        // PDF hint
        console.log(' KULLANIM İPUÇLARI:');
        console.log(' Ctrl+E: Düzenleme modunu aç/kapat');
        console.log(' Ctrl+S: Kaydet');
        console.log(' HTML İndir: Raporu HTML dosyası olarak indir');
        console.log(' Sıfırla: Tüm değişiklikleri geri al');
    </script>
</body>
</html>
"""
        return _translate_html_static_labels(html, lang_code)

    def _html_executive_summary(self, es: Dict, root_causes: List[Dict]) -> str:
        """Yönetici özeti HTML."""
        html = f"""
        <div class="section" id="executive-summary">
            <div class="section-header">1. {_L('section_executive_summary')}</div>
            
            <div class="subsection-header">1.1 {_L('subsection_incident_summary')}</div>
"""
        
        for field in ["what_happened", "immediate_response"]:
            if es.get(field):
                html += f'<div class="paragraph" contenteditable="true">{strip_hse_codes(str(es[field]))}</div>\n'
        
        html += f"""
            <div class="subsection-header">1.2 {_L('subsection_key_findings')}</div>
            <ul class="bullet-list">
"""
        for finding in es.get("key_findings", []):
            html += f'<li contenteditable="true">{strip_hse_codes(str(finding))}</li>\n'
        
        html += f"""
            </ul>
            
            <div class="subsection-header">1.3 {_L('subsection_immediate_actions')}</div>
            <table>
                <thead>
                    <tr>
                        <th>{_L('th_immediate_action')}</th>
                        <th>{_L('th_responsible')}</th>
                        <th>{_L('th_status')}</th>
                    </tr>
                </thead>
                <tbody>
"""
        
        for act in es.get("immediate_actions", []):
            html += f"""
                    <tr>
                        <td contenteditable="true">{act.get('action', '')}</td>
                        <td contenteditable="true">{act.get('responsible', '')}</td>
                        <td contenteditable="true">{act.get('status', '')}</td>
                    </tr>
"""
        
        html += """
                </tbody>
            </table>
        </div>
"""
        return html

    def _html_incident_details(self, details: Dict) -> str:
        """Olay bilgileri HTML."""
        html = f"""
        <div class="section" id="incident-details">
            <div class="section-header">2. {_L('section_incident_details')}</div>
            
            <div class="subsection-header">2.1 {_L('subsection_info_table')}</div>
            <table>
"""
        
        for key, val in details.get("info_table", {}).items():
            html += f"""
                <tr>
                    <td style="background: #D6E4F0; font-weight: bold; color: #1B3A5C;">{key}</td>
                    <td contenteditable="true">{val}</td>
                </tr>
"""
        
        html += f"""
            </table>
            
            <div class="subsection-header">2.2 {_L('subsection_event_details')}</div>
"""
        
        summary_only = strip_hse_codes(str(details.get("event_summary_only") or ""))
        if summary_only:
            html += f"""
            <div class="event-summary-compact">
                <span class="label">{_L('event_summary_label')}</span>
                <span contenteditable="true">{summary_only}</span>
            </div>
            <table>
"""
        else:
            html += """
            <table>
"""
        
        for key, val in details.get("event_table", {}).items():
            html += f"""
                <tr>
                    <td style="background: #D6E4F0; font-weight: bold; color: #1B3A5C;">{key}</td>
                    <td contenteditable="true">{val}</td>
                </tr>
"""
        
        html += f"""
            </table>
            
            <div class="subsection-header">2.3 {_L('subsection_timeline')}</div>
            <div class="timeline">
"""
        
        for step in details.get("timeline", []):
            html += f"""
                <div class="timeline-item">
                    <div class="timeline-time" contenteditable="true">{step.get('time', '')}</div>
                    <div class="timeline-event" contenteditable="true">{step.get('event', '')}</div>
                </div>
"""
        
        html += """
            </div>
        </div>
"""
        return html

    def _html_analysis_method(self, method: Dict) -> str:
        """Analiz yöntemi HTML."""
        method = _normalize_analysis_method(method)
        html = f"""
        <div class="section" id="analysis-method">
            <div class="section-header">3. {_L('section_analysis_method')}</div>
            
            <div class="subsection-header">3.1 {_L('subsection_five_why')}</div>
            <div class="paragraph" contenteditable="true">{method.get('five_why_explanation', '')}</div>
            
            <div class="subsection-header">3.2 {_L('subsection_code_system')}</div>
            <table>
                <thead>
                    <tr>
                        <th>{_L('th_code')}</th>
                        <th>{_L('th_category')}</th>
                        <th>{_L('th_description')}</th>
                    </tr>
                </thead>
                <tbody>
"""
        
        for code in method.get("code_system", []):
            html += f"""
                    <tr>
                        <td style="font-weight: bold; color: #1B3A5C;">{code.get('code', '')}</td>
                        <td contenteditable="true">{code.get('category', '')}</td>
                        <td contenteditable="true">{code.get('description', '')}</td>
                    </tr>
"""
        
        html += f"""
                </tbody>
            </table>
            
            <div class="subsection-header">3.3 {_L('subsection_analysis_team')}</div>
            <table>
                <thead>
                    <tr>
                        <th>{_L('th_name')}</th>
                        <th>{_L('th_role')}</th>
                        <th>{_L('th_date')}</th>
                    </tr>
                </thead>
                <tbody>
"""
        
        for member in method.get("team_members", []):
            html += f"""
                    <tr>
                        <td contenteditable="true">{member.get('name', '')}</td>
                        <td contenteditable="true">{member.get('role', '')}</td>
                        <td contenteditable="true">{member.get('date', '')}</td>
                    </tr>
"""
        
        html += """
                </tbody>
            </table>
        </div>
"""
        return html

    def _html_branches(self, branches: List[Dict]) -> str:
        """Analiz dalları HTML."""
        html = """
        <div class="section" id="branches">
"""
        
        for branch in branches:
            bn = branch.get("branch_number", 1)
            html += f"""
            <div class="subsection">
            <div class="section-header">{3+bn}. {branch.get('branch_title', f"{_L('branch_critical_factor')} {bn}")}</div>
            
            <div class="subsection-header">{3+bn}.1 {_L('subsection_why_chain')}</div>
            <div class="why-chain">
"""
            
            why_chain = branch.get("why_chain", [])
            # P1.23-G4: BARSEL kod rozeti — varsayılan KAPALI. Mevcut "kodları gizle"
            # tasarımı korunur; yalnızca REPORT_SHOW_WHY_CODES=1 ile yapısal rozet eklenir.
            show_why_codes = (
                os.getenv("REPORT_SHOW_WHY_CODES", "0").strip().lower()
                in ("1", "true", "yes", "on")
            )
            for idx, why in enumerate(why_chain):
                qtxt = strip_hse_codes(str(why.get('question', '') or why.get('question_tr', '') or ''))
                atxt = strip_hse_codes(str(why.get('answer', '') or why.get('answer_tr', '') or ''))
                wn = why.get('number') or why.get('level') or (idx + 1)
                code_badge = ""
                if show_why_codes:
                    why_code = str(why.get('code', '') or why.get('hsg245_code', '') or '').strip().upper()
                    if why_code:
                        code_badge = f'<span class="why-code">{why_code}</span> '
                html += f"""
                <div class="why-item">
                    <div class="why-number">{_L('why_prefix')} {wn}</div>
                    <div class="why-question" contenteditable="true">{qtxt}</div>
                    <div class="why-answer" contenteditable="true">{code_badge}→ {format_report_html_rich(atxt)}</div>
                </div>
"""
            
            colors = ['red', 'orange', 'green', 'blue']
            color = colors[(bn - 1) % len(colors)]
            
            root_cause_title = strip_root_cause_label_prefix(
                str(branch.get("root_cause_title", "") or ""),
                branch_number=bn,
            )
            
            html += f"""
            </div>
            
            <div class="subsection-header">{3+bn}.2 {_L('subsection_root_cause')}</div>
            <div class="colored-box box-{color}">
                <div class="box-header" contenteditable="true">{_L('root_cause_prefix')} {bn}: {root_cause_title}</div>
                <div class="box-content" contenteditable="true">{format_report_html_rich(str(branch.get('root_cause_detail', '') or ''))}</div>
            </div>
"""
            
            if branch.get("organizational_factors"):
                html += f"""
            <div class="subsection-header">{3+bn}.3 {_L('subsection_org_factors')}</div>
            <ul class="bullet-list">
"""
                for factor in branch.get("organizational_factors", []):
                    html += f'<li contenteditable="true">{strip_hse_codes(str(factor))}</li>\n'
                html += """
            </ul>
"""
            
            html += """
            </div>
"""
        
        html += """
        </div>
"""
        
        return html

    def _ensure_root_causes_from_branches(
        self,
        root_causes: List[Dict[str, Any]],
        branches: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        """Rapor kök nedenleri eksikse dallardan tamamla (5 dal -> 5 kök neden)."""
        existing: List[Dict[str, Any]] = []
        seen: set[str] = set()
        for rc in root_causes or []:
            if not isinstance(rc, dict):
                continue
            title = strip_hse_codes(str(rc.get("title", "") or "")).strip()
            detail = strip_hse_codes(str(rc.get("detailed_description", "") or "")).strip()
            key = self._normalized_key(title or detail)
            if key and key not in seen:
                seen.add(key)
                existing.append(rc)

        for br in branches or []:
            if not isinstance(br, dict):
                continue
            title = strip_hse_codes(str(br.get("root_cause_title", "") or "")).strip()
            detail = strip_hse_codes(str(br.get("root_cause_detail", "") or "")).strip()
            if not title and not detail:
                continue
            key = self._normalized_key(title or detail)
            if key in seen:
                continue
            seen.add(key)
            existing.append(
                {
                    "title": title or "Kök Neden",
                    "category": "",
                    "contributing_organizations": "",
                    "detailed_description": detail or title or "",
                    "impacts": [],
                }
            )
        return existing

    @staticmethod
    def _normalized_key(text: str) -> str:
        s = (text or "").lower().strip()
        s = re.sub(r"[^\w\s]", " ", s)
        s = re.sub(r"\s+", " ", s)
        return s

    def _html_root_causes(self, root_causes: List[Dict], section_no: int = 6) -> str:
        """Nihai kök nedenler HTML."""
        html = """
        <div class="section" id="root-causes">
"""
        html += f"""            <div class="section-header">{section_no}. {_L('section_final_root_causes')}</div>
"""
        
        for i, rc in enumerate(root_causes):
            html += f"""
            <div class="root-cause-box root-cause-{i+1}">
                <div class="root-cause-header" contenteditable="true">{_L('root_cause_prefix')} {i+1}: {strip_root_cause_label_prefix(str(rc.get('title', '') or ''), branch_number=i + 1)}</div>
                <div class="root-cause-content">
                    <table style="margin-bottom: 20px;">
                        <tr>
                            <td style="background: #D6E4F0; font-weight: bold;">{_L('th_category')}</td>
                            <td contenteditable="true">{strip_hse_codes(str(rc.get('category', '') or ''))}</td>
                        </tr>
                        <tr>
                            <td style="background: #D6E4F0; font-weight: bold;">{_L('th_related_units')}</td>
                            <td contenteditable="true">{strip_hse_codes(str(rc.get('contributing_organizations', '') or ''))}</td>
                        </tr>
                    </table>
                    
                    <div class="paragraph" contenteditable="true">{strip_hse_codes(str(rc.get('detailed_description', '') or ''))}</div>
                    
                    <h4 style="margin-top: 20px; color: #1B3A5C;">{_L('impacts_from_cause')}</h4>
                    <ul class="bullet-list">
"""
            
            for impact in rc.get("impacts", []):
                html += f'<li contenteditable="true">{strip_hse_codes(str(impact))}</li>\n'
            
            html += """
                    </ul>
                </div>
            </div>
"""
        
        html += """
        </div>
"""
        return html

    def _html_meta_root_cause(self, meta: Dict) -> str:
        """Meta kök neden HTML (tüm dalların ortak paydası)."""
        if not meta or not meta.get("exists"):
            return ""
        
        title = strip_hse_codes(str(meta.get('title', 'Meta Kök Neden') or ''))
        description = strip_hse_codes(str(meta.get('description', '') or ''))
        
        html = f"""
        <div class="section" id="meta-root-cause" style="page-break-before: always;">
            <div class="subsection">
                <h3 style="color: #C0392B;">🎯 Stratejik Kök Neden (Tüm Dalların Ortak Paydası)</h3>
                <div class="alert-box" style="background: linear-gradient(to right, #C0392B, #E74C3C); color: white; padding: 20px; border-radius: 8px; margin: 15px 0;">
                    <h4 contenteditable="true" style="color: white; font-size: 18px; margin-top: 0;">
                        <strong>{title}</strong>
                    </h4>
                    <p contenteditable="true" style="font-size: 14px; line-height: 1.8;">
                        {description}
                    </p>
                </div>
            </div>
"""
        
        synthesized = meta.get('synthesized_from', [])
        if synthesized:
            html += f"""
            <div class="subsection">
                <h3 style="color: #C0392B;">🔗 Sentezlenen Kök Nedenler</h3>
                <p contenteditable="true" style="margin-bottom: 10px;">
                    Bu meta kök neden, analizdeki <strong>{len(synthesized)}</strong> dalın ortak paydasından türetilmiş üst düzey bir sistemik zayıflığı ifade eder.
                </p>
            </div>
"""
        
        # Sistemik zayıflık
        systemic = strip_hse_codes(str(meta.get('systemic_weakness', '') or ''))
        if systemic:
            html += f"""
            <div class="subsection">
                <h3 style="color: #C0392B;">⚠️ Sistemik Zayıflık</h3>
                <p contenteditable="true" style="background: #FFF3E0; padding: 15px; border-left: 4px solid #E67E22; border-radius: 4px;">
                    {systemic}
                </p>
            </div>
"""
        
        # Stratejik sonuçlar
        implications = meta.get('strategic_implications', [])
        if implications:
            html += """
            <div class="subsection">
                <h3 style="color: #C0392B;">📊 Stratejik Sonuçlar ve Etkiler</h3>
                <ul style="list-style: none; padding-left: 0;">
"""
            for i, imp in enumerate(implications, 1):
                imp_t = strip_hse_codes(str(imp))
                html += f"""
                    <li style="background: #FFE6E6; padding: 12px; margin: 8px 0; border-left: 4px solid #E74C3C; border-radius: 4px;">
                        <strong style="color: #C0392B;">{i}.</strong> <span contenteditable="true">{imp_t}</span>
                    </li>
"""
            html += """
                </ul>
            </div>
"""
        
        html += """
        </div>
"""
        return html

    def _html_contributing_factors(self, factors: List[Dict], section_no: int = 6) -> str:
        """Sistemsel faktörler HTML."""
        html = """
        <div class="section" id="contributing-factors">
"""
        html += f"""            <div class="section-header">{section_no}. {_L('section_systemic_factors')}</div>
            
            <table>
                <thead>
                    <tr>
                        <th>{_L('th_factor_type')}</th>
                        <th>{_L('th_description')}</th>
                        <th>{_L('th_impact_level')}</th>
                    </tr>
                </thead>
                <tbody>
"""
        
        for factor in factors:
            impact = factor.get("impact_level", "Orta")
            impact_class = f"impact-{impact.lower()}" if impact.lower() in ['high', 'yüksek'] else (f"impact-medium" if impact.lower() in ['medium', 'orta'] else "impact-low")
            
            html += f"""
                    <tr>
                        <td style="font-weight: bold;" contenteditable="true">{factor.get('factor_type', '')}</td>
                        <td contenteditable="true">{factor.get('description', '')}</td>
                        <td class="{impact_class}" contenteditable="true">{impact}</td>
                    </tr>
"""
        
        html += """
                </tbody>
            </table>
        </div>
"""
        return html

    def _html_corrective_actions(self, actions: List[Dict], section_no: int = 7) -> str:
        """Düzeltici faaliyetler HTML."""
        html = """
        <div class="section" id="corrective-actions">
"""
        html += f"""            <div class="section-header">{section_no}. {_L('section_corrective_actions')}</div>
            
            <table>
                <thead>
                    <tr>
                        <th style="width: 5%;">{_L('th_no')}</th>
                        <th style="width: 35%;">{_L('th_activity')}</th>
                        <th style="width: 10%;">{_L('th_priority')}</th>
                        <th style="width: 15%;">{_L('th_responsible')}</th>
                        <th style="width: 10%;">{_L('th_duration')}</th>
                        <th style="width: 25%;">{_L('th_kpi')}</th>
                    </tr>
                </thead>
                <tbody>
"""
        
        for act in actions:
            priority = act.get("priority", "ORTA")
            priority_class = "priority-urgent" if priority == "ACİL" else ("priority-high" if priority == "YÜKSEK" else ("priority-medium" if priority == "ORTA" else "priority-low"))
            
            html += f"""
                    <tr>
                        <td>{act.get('no', '')}</td>
                        <td contenteditable="true">{act.get('action', '')}</td>
                        <td><span class="{priority_class}">{priority}</span></td>
                        <td contenteditable="true">{act.get('responsible', '')}</td>
                        <td contenteditable="true">{act.get('deadline', '')}</td>
                        <td contenteditable="true">{act.get('kpi', '')}</td>
                    </tr>
"""
        
        html += """
                </tbody>
            </table>
        </div>
"""
        return html

    def _html_lessons_learned(self, lessons: Dict, section_no: int = 8) -> str:
        """Çıkarılan dersler HTML."""
        sections = [
            (_L("lesson_what_to_do"), lessons.get("what_to_do", []), "green"),
            (_L("lesson_long_term"), lessons.get("long_term", []), "blue"),
            (_L("lesson_communication"), lessons.get("communication", []), "orange"),
            (_L("lesson_training"), lessons.get("training", []), "red"),
        ]
        
        html = """
        <div class="section" id="lessons-learned">
"""
        html += f"""            <div class="section-header">{section_no}. {_L('section_lessons_learned')}</div>
"""
        
        for title, items, color in sections:
            if items:
                content = "\n".join(f"• {item}" for item in items)
                html += f"""
            <div class="colored-box box-{color}">
                <div class="box-header">{title}</div>
                <div class="box-content" contenteditable="true">{content}</div>
            </div>
"""
        
        html += """
        </div>
"""
        return html

    def _html_conclusion(self, conclusion: Dict, section_no: int = 10) -> str:
        """Sonuç ve öneriler HTML."""
        html = f"""
        <div class="section" id="conclusion">
            <div class="section-header">{section_no}. {_L('section_conclusion')}</div>
            
            <div class="subsection-header">{section_no}.1 {_L('subsection_general_assessment')}</div>
            <div class="paragraph" contenteditable="true">{conclusion.get('overall_assessment', '')}</div>
            
            <div class="subsection-header">{section_no}.2 {_L('subsection_short_term')}</div>
            <ul class="bullet-list">
"""
        
        for measure in conclusion.get("short_term_measures", []):
            html += f'<li contenteditable="true">{measure}</li>\n'
        
        html += f"""
            </ul>
            
            <div class="subsection-header">{section_no}.3 {_L('subsection_long_term')}</div>
            <ul class="bullet-list">
"""
        
        for improvement in conclusion.get("long_term_improvements", []):
            html += f'<li contenteditable="true">{improvement}</li>\n'
        
        html += f"""
            </ul>
            
            <div class="subsection-header">{section_no}.4 {_L('subsection_comparison')}</div>
            <table class="comparison-table">
                <thead>
                    <tr>
                        <th>{_L('th_criterion')}</th>
                        <th>{_L('th_current')}</th>
                        <th>{_L('th_target')}</th>
                    </tr>
                </thead>
                <tbody>
"""
        
        for row in conclusion.get("comparison_table", []):
            html += f"""
                    <tr>
                        <td>{row.get('criterion', '')}</td>
                        <td class="current" contenteditable="true">{row.get('current', '')}</td>
                        <td class="target" contenteditable="true">{row.get('target', '')}</td>
                    </tr>
"""
        
        html += """
                </tbody>
            </table>
        </div>
"""
        return html

    def _html_signatures(self, section_no: int = 11) -> str:
        """İmza sayfası HTML."""
        html = """
        <div class="section signature-section" id="signatures">
"""
        html += f"""            <div class="section-header">{section_no}. {_L('section_signatures')}</div>
            
            <table class="signature-table">
                <thead>
                    <tr>
                        <th>{_L('th_role')}</th>
                        <th>{_L('th_name')}</th>
                        <th>{_L('th_title')}</th>
                        <th>{_L('th_signature')}</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td style="font-weight: bold; color: #1B3A5C;">{_L('sig_role_prepared')}</td>
                        <td contenteditable="true">HSE Uzmanı</td>
                        <td contenteditable="true">HSE Kök Neden Analisti</td>
                        <td contenteditable="true">
                            <div class="signature-line">
                                _____________________<br>
                                _____ / _____ / _____
                            </div>
                        </td>
                    </tr>
                    <tr>
                        <td style="font-weight: bold; color: #1B3A5C;">{_L('sig_role_reviewed')}</td>
                        <td contenteditable="true">HSE Yöneticisi</td>
                        <td contenteditable="true">HSE Departman Yöneticisi</td>
                        <td contenteditable="true">
                            <div class="signature-line">
                                _____________________<br>
                                _____ / _____ / _____
                            </div>
                        </td>
                    </tr>
                    <tr>
                        <td style="font-weight: bold; color: #1B3A5C;">{_L('sig_role_approved')}</td>
                        <td contenteditable="true">Tesis Müdürü</td>
                        <td contenteditable="true">Genel Operasyon Müdürü</td>
                        <td contenteditable="true">
                            <div class="signature-line">
                                _____________________<br>
                                _____ / _____ / _____
                            </div>
                        </td>
                    </tr>
                </tbody>
            </table>
            
            <div style="margin-top: 40px; padding: 20px; background: #F5F5F5; border-left: 4px solid #2E6DA4;">
                <p style="margin: 0; color: #666;">
                    <strong> Note:</strong> {_L('html_edit_note')}
                </p>
            </div>
        </div>
"""
        return html

    def _html_decision_tree(self, investigation_data: Dict, section_no: int = 12) -> str:
        """5-Why Decision Tree bölümü — Mermaid diagram embedded."""
        from agents.decision_tree_mermaid import DecisionTreeGenerator
        
        try:
            # RCA verisini çıkar
            rca_data = None
            incident_title = "Kaza Analizi"
            
            if "part3_rca" in investigation_data:
                rca_data = investigation_data["part3_rca"]
                # Olay başlığını part1'den al
                if "part1" in investigation_data and "overview" in investigation_data["part1"]:
                    overview = investigation_data["part1"]["overview"]
                    incident_title = overview.get("what_happened", "Kaza Analizi")[:100]
            elif "analysis_branches" in investigation_data:
                rca_data = investigation_data
                incident_title = rca_data.get("incident_event", "Kaza Analizi")
            
            if not rca_data or not rca_data.get("analysis_branches", rca_data.get("branches")):
                return ""  # Decision tree için veri yok
            
            # Decision tree generator
            gen = DecisionTreeGenerator()
            mermaid_code = gen._generate_mermaid_graph(
                rca_data.get("analysis_branches", rca_data.get("branches", [])),
                incident_title
            )
            
            html = f"""
        <div class="section" id="decision-tree" style="page-break-before: always;">
            <div class="section-header">{section_no}. 5-WHY KARAR AĞACI (DECISION TREE)</div>
            <p style="font-size: 13px; color: #555; margin: 8px 0 12px;">
                Üstten alta: OLAY → soru (kesik çerçeve) → cevap → kök neden. Yazdırırken dikey A4 için uygundur.
            </p>
            <div id="decision-tree-diagram" style="background: white; padding: 12px; border: 1px solid #ddd; width: 100%; min-height: min(95vh, 1400px); overflow: auto;">
                <div class="mermaid" style="width: 100%; min-height: 90vh;">
{mermaid_code}
                </div>
            </div>
            
            <!-- Mermaid.js library -->
            <script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
            <script>
                if (typeof mermaid !== 'undefined') {{
                    mermaid.initialize({{
                        startOnLoad: true,
                        theme: 'default',
                        flowchart: {{
                            useMaxWidth: false,
                            htmlLabels: true,
                            curve: 'basis',
                            padding: 20,
                            nodeSpacing: 28,
                            rankSpacing: 56,
                            diagramPadding: 16
                        }},
                        themeVariables: {{
                            fontSize: '14px',
                            fontFamily: 'Segoe UI, Arial, sans-serif',
                            primaryColor: '#fff',
                            primaryTextColor: '#222',
                            primaryBorderColor: '#333',
                            lineColor: '#333',
                            secondaryColor: '#f5f5f5',
                            tertiaryColor: '#fff'
                        }}
                    }});
                }} else {{
                    setTimeout(function() {{
                        if (typeof mermaid !== 'undefined') {{
                            mermaid.initialize({{
                                startOnLoad: true,
                                theme: 'default',
                                flowchart: {{
                                    useMaxWidth: false,
                                    htmlLabels: true,
                                    curve: 'basis',
                                    padding: 20,
                                    nodeSpacing: 28,
                                    rankSpacing: 56,
                                    diagramPadding: 16
                                }},
                                themeVariables: {{
                                    fontSize: '14px',
                                    fontFamily: 'Segoe UI, Arial, sans-serif'
                                }}
                            }});
                            mermaid.contentLoaded();
                        }}
                    }}, 1000);
                }}
            </script>
        </div>
"""
            return html
            
        except Exception as e:
            print(f"  Uyarı: Decision tree oluşturulamadı: {str(e)}")
            return ""

    def _html_incident_photos(self) -> str:
        """Olay fotoğrafları bölümü — 2 sayfa × 4 foto, tıkla-yükle."""
        pages_html = ""
        for page_no in range(1, 3):
            slots_html = ""
            for slot in range(1, 5):
                uid = f"photo_p{page_no}_s{slot}"
                slots_html += f"""
                <div class="photo-slot" id="{uid}_slot" onclick="triggerUpload('{uid}')">
                    <input type="file" id="{uid}_input" accept="image/*"
                           style="display:none"
                           onchange="loadPhoto(event, '{uid}')">
                    <button class="photo-remove" title="Fotoğrafı Kaldır"
                            onclick="event.stopPropagation(); removePhoto('{uid}')"></button>
                    <div class="upload-hint" id="{uid}_hint">
                        <span class="icon"></span>
                        <span class="label">Fotoğraf {slot}</span>
                        <span class="sub">Tıklayın veya sürükleyin</span>
                    </div>
                    <img id="{uid}_img" src="" alt="" style="display:none">
                    <input class="photo-caption"
                           id="{uid}_caption"
                           type="text"
                           placeholder="Fotoğraf açıklaması giriniz..."
                           onclick="event.stopPropagation()"
                           style="display:none">
                </div>"""

            pages_html += f"""
            <div class="photo-page" id="photo-page-{page_no}">
                <div class="photo-page-title">Sayfa {page_no} / 2</div>
                <div class="photo-grid">
                    {slots_html}
                </div>
            </div>"""

        return f"""
        <div class="section photo-section" id="incident-photos">
            <div class="section-header"> OLAY FOTOĞRAFLARI</div>
            <p style="color:#555; margin-bottom:20px;">
                Her kareye tıklayarak olay fotoğrafı yükleyin. Yüklenen fotoğraflar yalnızca 
                bu rapor oturumunda saklanır; raporu <strong>HTML olarak indirerek</strong> 
                kalıcı hale getirebilirsiniz.
            </p>
            {pages_html}
        </div>

        <script>
        /*  OLAY FOTOĞRAFLARI JS  */
        function triggerUpload(uid) {{
            document.getElementById(uid + '_input').click();
        }}

        function loadPhoto(event, uid) {{
            const file = event.target.files[0];
            if (!file) return;
            const reader = new FileReader();
            reader.onload = function(e) {{
                const slot  = document.getElementById(uid + '_slot');
                const img   = document.getElementById(uid + '_img');
                const hint  = document.getElementById(uid + '_hint');
                const cap   = document.getElementById(uid + '_caption');
                img.src = e.target.result;
                img.style.display = 'block';
                hint.style.display = 'none';
                cap.style.display  = 'block';
                slot.classList.add('has-photo');
                // Auto-fill caption with filename (without extension)
                if (!cap.value) {{
                    cap.value = file.name.replace(/\\.[^/.]+$/, '');
                }}
            }};
            reader.readAsDataURL(file);
        }}

        function removePhoto(uid) {{
            const slot  = document.getElementById(uid + '_slot');
            const img   = document.getElementById(uid + '_img');
            const hint  = document.getElementById(uid + '_hint');
            const cap   = document.getElementById(uid + '_caption');
            const input = document.getElementById(uid + '_input');
            img.src = '';
            img.style.display = 'none';
            hint.style.display = 'flex';
            cap.style.display  = 'none';
            cap.value = '';
            input.value = '';
            slot.classList.remove('has-photo');
        }}

        // Sürükle-bırak desteği
        document.querySelectorAll('.photo-slot').forEach(function(slot) {{
            slot.addEventListener('dragover', function(e) {{
                e.preventDefault();
                slot.style.borderColor = '#1B3A5C';
                slot.style.background  = '#D6E4F0';
            }});
            slot.addEventListener('dragleave', function() {{
                if (!slot.classList.contains('has-photo')) {{
                    slot.style.borderColor = '#2E6DA4';
                    slot.style.background  = '#F8FBFF';
                }}
            }});
            slot.addEventListener('drop', function(e) {{
                e.preventDefault();
                slot.style.borderColor = '';
                slot.style.background  = '';
                const uid   = slot.id.replace('_slot', '');
                const input = document.getElementById(uid + '_input');
                const file  = e.dataTransfer.files[0];
                if (file && file.type.startsWith('image/')) {{
                    // Simulate file input change
                    const dt = new DataTransfer();
                    dt.items.add(file);
                    input.files = dt.files;
                    loadPhoto({{target: input}}, uid);
                }}
            }});
        }});
        </script>
"""


if __name__ == "__main__":
    print("=" * 70)
    print(" SkillBasedDocxAgent V2 — Standalone Test")
    print("=" * 70)

    outputs = sorted(Path("outputs").glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
    if outputs:
        json_file = outputs[0]
        print(f" Kullanılan veri: {json_file}")
        with open(json_file, encoding="utf-8") as f:
            data = json.load(f)
    else:
        print(" outputs/*.json bulunamadı!")
        sys.exit(1)

    agent = SkillBasedDocxAgent()
    try:
        out = agent.generate_report(
            investigation_data=data,
            output_path="outputs/HSE_FULL_REPORT_V2.docx",
        )
        print(f"\n BAŞARILI! → {out}")
    except Exception as e:
        import traceback
        print(f"\n HATA: {e}")
        traceback.print_exc()
        sys.exit(1)
